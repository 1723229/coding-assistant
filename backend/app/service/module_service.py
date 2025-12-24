"""
Module Service

业务逻辑层 - Module operations with tree structure support
POINT type modules create sessions and manage workspaces
"""
import os
import uuid
import logging
import json
import asyncio
import tempfile
from typing import Optional, AsyncGenerator
from pathlib import Path
from fastapi import Query, UploadFile

from app.api.chat_router import module_repo
from app.config.logging_config import log_print
from app.config import get_settings
from app.config.settings import FrameworkDatabaseConfig
from app.utils.model.response_model import BaseResponse, ListResponse
from app.db.repository import ModuleRepository, ProjectRepository, VersionRepository, SessionRepository, \
    MessageRepository
from app.db.schemas import ModuleCreate, ModuleUpdate, ModuleResponse, VersionCreate, VersionUpdate
from app.db.models.module import ModuleType, ContentStatus
from app.db.models.version import VersionStatus
from app.core.executor import get_sandbox_executor
from app.core.github_service import GitHubService
from app.utils.mysql_util import MySQLUtil
from datetime import datetime
from app.utils.prompt.prompt_build import generate_code_from_spec
from app.config.settings import ContainerConfig
from app.core.agent_service import AgentService

logger = logging.getLogger(__name__)
settings = get_settings()
message_repo = MessageRepository()


class ModuleService:
    """
    模块服务类

    提供模块相关的所有业务逻辑操作，支持树形结构
    POINT 类型创建 session 和 workspace，拉取代码
    NODE 类型只做功能说明，URL 与子节点共享
    """

    def __init__(self):
        self.sort_code = 10000
        self.module_repo = ModuleRepository()
        self.project_repo = ProjectRepository()
        self.version_repo = VersionRepository()
        self.session_repo = SessionRepository()
        self.agent_service = AgentService()
        self.db = MySQLUtil(
            host=FrameworkDatabaseConfig.HOST,
            port=FrameworkDatabaseConfig.PORT,
            user=FrameworkDatabaseConfig.USER,
            password=FrameworkDatabaseConfig.PASSWORD,
            database=FrameworkDatabaseConfig.DATABASE
        )

    def _find_prd_gen_dir(self, workspace_dir: Path) -> Path:
        """
        查找不区分大小写的 PRD-GEN 目录

        Args:
            workspace_dir: workspace 根目录

        Returns:
            PRD-GEN 目录的 Path 对象（如果不存在则返回默认路径）
        """
        docs_dir = workspace_dir / "docs"

        if docs_dir.exists() and docs_dir.is_dir():
            for item in docs_dir.iterdir():
                if item.is_dir() and item.name.upper() == "PRD-GEN":
                    return item

        # 回退到默认路径
        return workspace_dir / "docs" / "PRD-GEN"

    async def _check_container_limit(self) -> tuple[bool, str]:
        """
        检查容器数量是否达到阈值

        Returns:
            (是否可以创建, 错误消息)
        """
        try:
            running_count = await self.version_repo.count_running_containers()
            max_containers = ContainerConfig.MAX_RUNNING_CONTAINERS

            if running_count >= max_containers:
                return False, f"已达到最大容器数量限制({max_containers})，当前运行中: {running_count}"

            return True, ""

        except Exception as e:
            logger.error(f"Failed to check container limit: {e}", exc_info=True)
            return False, f"检查容器限制失败: {str(e)}"

    async def _update_version_status(
        self,
        version_id: int,
        status: VersionStatus,
        spec_content: Optional[str] = None
    ) -> bool:
        """
        更新 Version 状态

        Args:
            version_id: Version ID
            status: 新状态
            spec_content: Spec 内容（可选）

        Returns:
            是否更新成功
        """
        try:
            update_data = VersionUpdate(status=status.value)
            if spec_content:
                update_data.spec_content = spec_content

            updated = await self.version_repo.update_version(
                version_id=version_id,
                data=update_data
            )

            if updated:
                logger.info(f"Updated version {version_id} status to {status.value}")
                return True
            else:
                logger.error(f"Failed to update version {version_id} status")
                return False

        except Exception as e:
            logger.error(f"Failed to update version status: {e}", exc_info=True)
            return False

    async def _cleanup_container(self, module_id: int, session_id: str) -> bool:
        """
        清理容器

        Args:
            module_id: 模块ID
            session_id: 会话ID

        Returns:
            是否清理成功
        """
        try:
            logger.info(f"Cleaning up container for module {module_id}, session {session_id}")

            # 获取 sandbox executor
            executor = get_sandbox_executor()

            # 停止并删除容器
            await executor.stop_container(session_id)

            logger.info(f"Container cleanup completed for session {session_id}")
            return True

        except Exception as e:
            logger.error(f"Failed to cleanup container: {e}", exc_info=True)
            return False

    async def _generate_spec_document(
        self,
        require_content: str,
        module_name: str,
        module_code: str,
        workspace_path: str,
        session_id: str
    ) -> Optional[dict]:
        """
        使用 Claude 生成技术规格文档

        Args:
            require_content: 功能需求描述
            module_name: 模块名称
            module_code: 模块代码
            workspace_path: 工作空间路径
            session_id: 会话ID

        Returns:
            生成的规格文档路径，失败返回 None
        """
        try:
            logger.info(f"Generating spec document for module: {module_code}")

            # # 创建 spec 目录
            # spec_dir = Path(workspace_path) / "spec"
            # spec_dir.mkdir(parents=True, exist_ok=True)

            # 构建提示词
#             prompt = f"""你是一位资深的软件架构师和技术文档专家。请根据以下功能需求，生成一份详细的技术规格文档（Technical Specification Document）。
#
# 功能需求：
# {require_content}
#
# 注意：
# - 文档应该足够详细，让AI能够根据文档直接生成代码
# - 使用Markdown格式
# - 代码示例使用代码块
# - 表格使用Markdown表格格式
# - 流程图使用Mermaid语法
#
# 请直接输出Markdown格式的文档内容，不要有任何额外的解释或说明。"""
#
#             # 获取沙箱服务实例
#             sandbox_service = await session_manager.get_service(
#                 session_id=session_id,
#                 workspace_path=workspace_path,
#             )
#
#             # 调用沙箱服务生成文档
#             logger.info("Calling sandbox service to generate spec document...")
#             messages = await sandbox_service.chat(prompt=prompt, session_id=session_id)
#
#             # 提取文本内容
#             spec_content = ""
#             for msg in messages:
#                 if msg.type == "text":
#                     spec_content += msg.content + "\n"
#
#             if not spec_content.strip():
#                 logger.error("Claude did not generate any content")
#                 return None
            return {"spec_content": require_content}

        except Exception as e:
            logger.error(f"Failed to generate spec document: {e}", exc_info=True)
            return None

    @log_print
    async def list_modules(
        self,
        project_id: int,
        skip: int = Query(0, ge=0, description="Number of records to skip"),
        limit: int = Query(100, ge=1, le=500, description="Maximum number of records"),
    ):
        """获取项目的模块列表"""
        try:
            # Verify project exists
            project = await self.project_repo.get_project_by_id(project_id=project_id)
            if not project:
                return BaseResponse.not_found(message=f"项目 ID {project_id} 不存在")

            modules = await self.module_repo.get_modules_by_project(
                project_id=project_id,
                skip=skip,
                limit=limit
            )
            total = await self.module_repo.count_modules(project_id=project_id)

            items = [ModuleResponse.model_validate(m) for m in modules]
            return ListResponse.success(items=items, total=total)
        except Exception as e:
            logger.error(f"获取模块列表失败: {e}", exc_info=True)
            return BaseResponse.error(message=f"获取模块列表失败: {str(e)}")

    @log_print
    async def get_module_tree(self, project_id: int):
        """获取项目的模块树形结构"""
        try:
            # Verify project exists
            project = await self.project_repo.get_project_by_id(project_id=project_id)
            if not project:
                return BaseResponse.not_found(message=f"项目 ID {project_id} 不存在")

            tree = await self.module_repo.build_module_tree(project_id=project_id)
            return BaseResponse.success(data=tree, message="获取模块树成功")
        except Exception as e:
            logger.error(f"获取模块树失败: {e}", exc_info=True)
            return BaseResponse.error(message=f"获取模块树失败: {str(e)}")

    @log_print
    async def create_module(self, data: ModuleCreate, created_by: Optional[str] = None):
        """
        创建新模块

        流程：
        - NODE 类型：直接创建数据库记录，URL 与子节点共享
        - POINT 类型：
          1. 检查项目代码唯一性
          2. 生成 session_id 和 workspace_path
          3. 创建数据库记录
          4. 创建 Docker 容器（可选）
          5. 克隆 GitHub 仓库到 workspace
        """
        try:
            # Verify project exists and get project info
            project = await self.project_repo.get_project_by_id(project_id=data.project_id)
            if not project:
                return BaseResponse.not_found(message=f"项目 ID {data.project_id} 不存在")

            # Check if module code already exists in this project
            existing = await self.module_repo.get_module_by_code(
                project_id=data.project_id,
                is_active = 1,
                code=data.code
            )
            if existing:
                return BaseResponse.error(
                    message=f"模块代码 '{data.code}' 在项目中已存在"
                )

            # Verify parent exists if specified
            if data.parent_id:
                parent = await self.module_repo.get_module_by_id(module_id=data.parent_id)
                if not parent:
                    return BaseResponse.not_found(message=f"父模块 ID {data.parent_id} 不存在")
                if parent.project_id != data.project_id:
                    return BaseResponse.error(message="父模块必须属于同一项目")

            # Prepare module data
            module_data = data.model_dump()
            url_parent_id = module_data.pop("url_parent_id")

            # Handle POINT type: create session and workspace
            if data.type == ModuleType.POINT:
                logger.info(f"Creating POINT type module: {data.code}")

                # 1. Generate session_id and workspace_path
                session_id = str(uuid.uuid4())
                workspace_path = str(settings.workspace_base_path / session_id)

                module_data.update({
                    "session_id": session_id,
                    "workspace_path": workspace_path,
                    "branch": data.branch or "main",
                    "is_active": 1,
                })

                logger.info(f"Module session_id: {session_id}, workspace: {workspace_path}")


                # 2. Clone GitHub repository
                if project.codebase:
                    try:
                        logger.info(f"Cloning repository: {project.codebase}")
                        service = GitHubService(token=project.token)
                        await service.clone_repo(
                            repo_url=project.codebase,
                            target_path=workspace_path,
                            branch=module_data.get("branch"),
                        )
                        logger.info(f"Repository cloned successfully to: {workspace_path}")
                        await service.create_branch(repo_path=workspace_path, branch_name= data.branch + '-' + session_id)
                        logger.info(f"Branch created successfully to: {workspace_path}")
                    except Exception as e:
                        # Mark module as inactive if clone fails
                        return BaseResponse.error(
                            message=f"代码拉取失败: {str(e)}"
                        )

                # 3. Create module in database first
                module = await self.module_repo.create_module(
                    data=module_data,
                    created_by=created_by
                )
                await self.session_repo.create_session(
                    session_id=session_id,
                    name=project.code + '-' + data.code,
                    workspace_path=workspace_path,
                    github_repo_url=project.codebase,
                    github_branch=data.branch or "main",
                )
                module_id = module.id
                logger.info(f"Module created in database: {module_id}")

                # 5. Generate spec document if require_content is provided
                commit_id = None

                if data.require_content:
                    try:
                        logger.info(f"Generating spec document for module: {module_id}")
                        spec_dict = await self._generate_spec_document(
                            require_content=data.require_content,
                            module_name=data.name,
                            module_code=data.code,
                            workspace_path=workspace_path,
                            session_id=session_id
                        )
                        await self.module_repo.update_module(module_id=module_id, data=module_data)
                        if spec_dict["spec_content"]:
                            logger.info(f"Spec document generated successfully")

                            # 6. Generate code from spec and commit
                            logger.info(f"Generating code from spec for module: {module_id}")
                            commit_id, _ = await generate_code_from_spec(
                                spec_content=data.require_content,
                                workspace_path=workspace_path,
                                session_id=session_id,
                                module_code=data.code,
                                module_name=data.name,
                                module_url=data.url,
                                task_type="spec"
                            )

                            if commit_id:
                                logger.info(f"Code generated and committed successfully: {commit_id}")
                                module_data.update({
                                    "latest_commit_id": commit_id
                                })
                                await self.module_repo.update_module(module_id=module_id, data=module_data)

                                # 7. Savecommit to version table
                                try:
                                    # Generate version code based on timestamp
                                    version_code = f"v{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                                    version_data = VersionCreate(
                                        code=version_code,
                                        module_id=module_id,
                                        msg=f"{data.name} ({data.code}) 功能实现",
                                        commit=commit_id
                                    )
                                    version = await self.version_repo.create_version(
                                        data=version_data,
                                        created_by=created_by
                                    )
                                    logger.info(f"Version created: {version.id}, code: {version_code}")
                                except Exception as e:
                                    logger.error(f"Failed to create version record: {e}", exc_info=True)
                            else:
                                logger.warning("Code generation did not produce a commit")
                        else:
                            # await self.module_repo.delete_module(module_id=module_id)
                            return BaseResponse.error(message=f"创建模块失败: Spec document generation returned None")

                    except Exception as e:
                        logger.error(f"Failed in spec/code generation process: {e}", exc_info=True)
                        # await self.module_repo.delete_module(module_id=module_id)
                        return BaseResponse.error(message=f"创建模块失败: {str(e)}")
                        # Continue even if spec/code generation fails

                # 8. Refresh module data and add commit_id
                module = await self.module_repo.get_module_by_id(module_id=module_id)
                module.spec_file_path = spec_dict.get("spec_file_path")
                module.spec_content = spec_dict.get("spec_content")
                response_data = ModuleResponse.model_validate(module)

                # 创建沙箱容器 (optional, non-blocking)
                try:
                    executor = get_sandbox_executor()
                    container_info = await executor.create_workspace(
                        session_id=session_id,
                        workspace_path=workspace_path
                    )
                    logger.info(f"Sandbox container created: {container_info['id']}")
                    module_data.update({
                        "container_id": container_info["id"],
                    })
                    await self.module_repo.update_module(module_id=module_id, data=module_data)
                except Exception as e:
                    logger.warning(f"Failed to create container for module {module_id}: {e}")
                    # await self.module_repo.delete_module(module_id=module_id)
                    return BaseResponse.error(message=f"创建模块失败: {str(e)}")

                    # Continue even if container creation fails
                menu = {"full_name": module.name, "english_name": module.code, "url_address": module.url, "enable_mark": 1, "parent_id": url_parent_id, "sort_code": self.sort_code}
                insert_id = self.db.insert(table='sys_module', data=menu)
                module_data.update({
                    "url_id": insert_id,
                    "preview_url": settings.preview_ip + ':' + str(container_info["code_port"]) + data.url,
                })
                logger.info(f"preview_url: {settings.preview_ip + ':' + str(container_info['code_port']) + data.url}")
                await self.module_repo.update_module(module_id=module_id, data=module_data)
                return BaseResponse.success(
                    data=response_data,
                    message="POINT 模块创建成功，代码已拉取并生成"
                )

            else:
                # NODE type: simple creation without workspace
                menu = {"full_name": data.name, "english_name": data.code, "url_address": '/',
                        "enable_mark": 1, "sort_code": self.sort_code}
                insert_id = self.db.insert(table='sys_module', data=menu)
                module_data.update({
                    "url_id": insert_id,
                })
                logger.info(f"Creating NODE type module: {data.code}")
                module = await self.module_repo.create_module(
                    data=module_data,
                    created_by=created_by
                )
                return BaseResponse.success(
                    data=ModuleResponse.model_validate(module),
                    message="NODE 模块创建成功"
                )

        except Exception as e:
            logger.error(f"创建模块失败: {e}", exc_info=True)
            # await self.module_repo.delete_module(module_id=module_id)
            # self.db.execute_update("DELETE FROM sys_module WHERE module_id=%s", module_id)
            return BaseResponse.error(message=f"创建模块失败: {str(e)}")

    async def create_module_stream(self, data: ModuleCreate, created_by: Optional[str] = None) -> AsyncGenerator[str, None]:
        """
        流式创建新模块（SSE）

        实时返回创建进度，适用于POINT类型模块的长时间操作

        流程：
        1. 验证项目和模块合法性
        2. 创建工作空间
        3. 模块入库和sys_module入库
        4. 检查并拉取代码（如果project有git地址）
        5. 检查容器阈值
        6. 创建容器
        7. 生成Spec（如果有require_content）

        生成的事件：
        - connected: 连接建立
        - step: {step: "step_name", status: "progress/success/error", message: "...", progress: 20}
        - complete: {module_id: xxx, session_id: xxx}
        - error: {message: "error message"}
        """
        module_id = None
        insert_id = None
        session_id = None
        workspace_path = None
        version_id = None

        try:
            # 发送连接确认
            yield f"data: {json.dumps({'type': 'connected'})}\n\n"
            await asyncio.sleep(0)

            # 步骤1: 验证项目
            yield f"data: {json.dumps({'type': 'step', 'step': 'validate_project', 'status': 'progress', 'message': '验证项目信息...', 'progress': 5})}\n\n"

            project = await self.project_repo.get_project_by_id(project_id=data.project_id)
            if not project:
                yield f"data: {json.dumps({'type': 'error', 'message': f'项目 ID {data.project_id} 不存在'})}\n\n"
                return

            yield f"data: {json.dumps({'type': 'step', 'step': 'validate_project', 'status': 'success', 'message': '项目验证成功', 'progress': 10})}\n\n"

            # 步骤2: 检查模块代码唯一性
            yield f"data: {json.dumps({'type': 'step', 'step': 'check_code', 'status': 'progress', 'message': '检查模块代码唯一性...', 'progress': 15})}\n\n"

            existing = await self.module_repo.get_module_by_code(
                project_id=data.project_id,
                is_active=1,
                code=data.code
            )
            if existing:
                yield f"data: {json.dumps({'type': 'error', 'message': f'模块代码 {data.code} 在项目中已存在'})}\n\n"
                return

            yield f"data: {json.dumps({'type': 'step', 'step': 'check_code', 'status': 'success', 'message': '模块代码检查通过', 'progress': 20})}\n\n"

            # 验证父节点
            if data.parent_id:
                parent = await self.module_repo.get_module_by_id(module_id=data.parent_id)
                if not parent or parent.project_id != data.project_id:
                    yield f"data: {json.dumps({'type': 'error', 'message': '父模块验证失败'})}\n\n"
                    return

            module_data = data.model_dump()
            url_parent_id = module_data.pop("url_parent_id")

            if data.type == ModuleType.POINT:
                # 步骤3: 生成session和创建工作空间
                yield f"data: {json.dumps({'type': 'step', 'step': 'create_workspace', 'status': 'progress', 'message': '创建工作空间...', 'progress': 25})}\n\n"

                session_id = str(uuid.uuid4())
                workspace_path = str(settings.workspace_base_path / session_id)

                # 创建工作空间目录
                workspace_dir = Path(workspace_path)
                workspace_dir.mkdir(parents=True, exist_ok=True)

                module_data.update({
                    "session_id": session_id,
                    "workspace_path": workspace_path,
                    "branch": data.branch or "main",
                    "is_active": 1,
                })

                yield f"data: {json.dumps({'type': 'step', 'step': 'create_workspace', 'status': 'success', 'message': f'工作空间创建成功: {workspace_path}', 'progress': 30})}\n\n"

                # 步骤4: 模块入库
                yield f"data: {json.dumps({'type': 'step', 'step': 'create_module', 'status': 'progress', 'message': '创建模块记录...', 'progress': 35})}\n\n"

                # 先插入 sys_module 表，获取 insert_id
                try:
                    menu = {
                        "full_name": data.name,
                        "english_name": data.code,
                        "url_address": data.url,
                        "enable_mark": 1,
                        "parent_id": url_parent_id,
                        "sort_code": self.sort_code
                    }
                    insert_id = self.db.insert(table='sys_module', data=menu)
                    logger.info(f"Inserted sys_module: {data.code}, url_id: {insert_id}")

                    # 将 url_id 添加到模块数据中
                    module_data["url_id"] = insert_id

                except Exception as e:
                    logger.error(f"Failed to insert sys_module: {e}")
                    yield f"data: {json.dumps({'type': 'error', 'message': f'sys_module入库失败: {str(e)}'})}\n\n"
                    return

                # 创建模块记录
                module = await self.module_repo.create_module(data=module_data, created_by=created_by)
                module_id = module.id

                yield f"data: {json.dumps({'type': 'step', 'step': 'create_module', 'status': 'success', 'message': f'模块ID: {module_id}, URL_ID: {insert_id}', 'module_id': module_id, 'progress': 40})}\n\n"

                # 步骤5: 检查并拉取代码
                if project.codebase and project.codebase != '':
                    if not project.token:
                        yield f"data: {json.dumps({'type': 'error', 'message': '项目配置了Git地址但缺少Token'})}\n\n"
                        return

                    # 检查工作空间是否已有代码
                    has_code = (workspace_dir / ".git").exists()

                    if not has_code:
                        yield f"data: {json.dumps({'type': 'step', 'step': 'clone_repo', 'status': 'progress', 'message': '正在克隆代码仓库...', 'progress': 45})}\n\n"

                        try:
                            service = GitHubService(token=project.token)
                            repo = await service.clone_repo(
                                repo_url=project.codebase,
                                target_path=workspace_path,
                                branch=module_data.get("branch"),
                            )
                            if repo:
                                yield f"data: {json.dumps({'type': 'step', 'step': 'clone_repo', 'status': 'success', 'message': '代码仓库克隆成功', 'progress': 50})}\n\n"
                            else:
                                yield f"data: {json.dumps({'type': 'error', 'message': f'代码拉取失败'})}\n\n"
                                return

                            await self.session_repo.create_session(
                                session_id=session_id,
                                name=project.code + '-' + data.code,
                                workspace_path=workspace_path,
                                github_repo_url=project.codebase,
                                github_branch=data.branch or "main",
                            )

                            # 创建分支
                            yield f"data: {json.dumps({'type': 'step', 'step': 'create_branch', 'status': 'progress', 'message': '创建功能分支...', 'progress': 55})}\n\n"

                            branch_name = f"{data.branch or 'main'}-{session_id}"
                            await service.create_branch(repo_path=workspace_path, branch_name=branch_name)

                            yield f"data: {json.dumps({'type': 'step', 'step': 'create_branch', 'status': 'success', 'message': '功能分支创建成功', 'progress': 60})}\n\n"

                        except Exception as e:
                            logger.error(f"Code clone failed: {e}")
                            yield f"data: {json.dumps({'type': 'error', 'message': f'代码拉取失败: {str(e)}'})}\n\n"
                            return
                    else:
                        yield f"data: {json.dumps({'type': 'step', 'step': 'clone_repo', 'status': 'success', 'message': '工作空间已有代码，跳过克隆', 'progress': 60})}\n\n"
                else:
                    yield f"data: {json.dumps({'type': 'error', 'message': '项目未配置Git地址，请配置好仓库信息和token', 'progress': 60})}\n\n"
                    return

                # 步骤6: 检查容器阈值
                yield f"data: {json.dumps({'type': 'step', 'step': 'check_container_limit', 'status': 'progress', 'message': '检查容器限制...', 'progress': 62})}\n\n"

                can_create, error_msg = await self._check_container_limit()
                if not can_create:
                    yield f"data: {json.dumps({'type': 'error', 'message': error_msg})}\n\n"
                    return
                yield f"data: {json.dumps({'type': 'step', 'step': 'check_container_limit', 'status': 'success', 'message': '容器限制检查通过', 'progress': 65})}\n\n"

                # 步骤7: 创建沙箱容器
                yield f"data: {json.dumps({'type': 'step', 'step': 'create_container', 'status': 'progress', 'message': '创建沙箱容器...', 'progress': 67})}\n\n"

                try:
                    executor = get_sandbox_executor()
                    container_info = await executor.create_workspace(
                        session_id=session_id,
                        workspace_path=workspace_path
                    )
                    module_data.update({"container_id": container_info["id"]})
                    container_id_short = container_info["id"][:12]
                    yield f'data: {json.dumps({"type": "step", "step": "create_container", "status": "success", "message": f"容器ID: {container_id_short}", "preview_url": settings.preview_ip + ":" + str(container_info["code_port"]) + data.url, "progress": 65})}\n\n'
                except Exception as e:
                    logger.warning(f"Container creation failed: {e}")
                    yield f"data: {json.dumps({'type': 'error', 'message': f'容器创建失败: {str(e)}'})}\n\n"
                    return

                module_data.update({
                    "url_id": insert_id,
                    "preview_url": settings.preview_ip + ':' + str(container_info["code_port"]) + data.url,
                })
                logger.info(f"preview_url: {settings.preview_ip + ':' + str(container_info['code_port']) + data.url}")
                await self.module_repo.update_module(module_id=module_id, data=module_data)


                yield f"data: {json.dumps({'type': 'step', 'step': 'create_db_record', 'status': 'success', 'message': f'模块ID: {module_id}', 'module_id': module_id, 'progress': 75})}\n\n"


                # 创建 Version 记录（初始状态：SPEC_GENERATING）
                yield f"data: {json.dumps({'type': 'step', 'step': 'create_version', 'status': 'progress', 'message': '创建版本记录...', 'progress': 78})}\n\n"

                version_code = f"v1.0.0-{datetime.now().strftime('%Y%m%d%H%M%S')}"
                version_data = VersionCreate(
                    code=version_code,
                    module_id=module_id,
                    msg="[SpecCoding Auto Commit] - Initial spec generation",
                    status=VersionStatus.SPEC_GENERATING.value
                )

                version = await self.version_repo.create_version(data=version_data, created_by=created_by)
                version_id = version.id

                yield f"data: {json.dumps({'type': 'step', 'step': 'create_version', 'status': 'success', 'message': f'版本ID: {version_id}', 'version_id': version_id, 'progress': 79})}\n\n"

                # 步骤9: 生成spec文档
                if data.require_content:
                    module_update = ModuleUpdate(require_content=data.require_content, content_status=ContentStatus.COMPLETED)
                    module_repo.update_module(module_id=module.session_id, module_update=module_update)
                    try:
                        await message_repo.create_message(
                            session_id=session_id,
                            role='user',
                            content=data.require_content,
                            tool_name=None,
                            tool_input=None,
                            tool_result=None,
                        )
                    except Exception as e:
                        logger.error(f"Failed to save messages: {e}", exc_info=True)

                    yield f"data: {json.dumps({'type': 'step', 'step': 'generate_spec', 'status': 'progress', 'message': '正在生成spec文档...', 'progress': 80})}\n\n"

                    message_queue = asyncio.Queue()
                    try:
                        task = asyncio.create_task(generate_code_from_spec(
                            spec_content=data.require_content,
                            workspace_path=workspace_path,
                            session_id=session_id,
                            module_code=data.code,
                            module_name=data.name,
                            module_url=data.url,
                            task_type="spec",
                            message_queue=message_queue,)
                        )
                        buffer = ""
                        while True:
                            chunk = await message_queue.get()
                            if chunk is None:
                                # 最后 flush 剩余内容（即使没有 \n）
                                if buffer:
                                    yield f"data: {json.dumps({'type': 'step', 'step': 'ai_think', 'status': 'success', 'message': 'ai思考...', 'ai_message': buffer, 'progress': 85}, ensure_ascii=False)}\n\n"
                                break
                            # 累加到缓冲区
                            buffer += chunk.content
                            # 持续检查缓冲区中是否有完整的行（以 \n 结尾）
                            while "\n\n" in buffer:
                                # 分割出第一个完整行（包含 \n）
                                line, buffer = buffer.split("\n\n", 1)
                                # 推送这一整行
                                yield f"data: {json.dumps({'type': 'step', 'step': 'ai_think', 'status': 'success', 'message': 'ai思考...', 'ai_message': line, 'progress': 85}, ensure_ascii=False)}\n\n"
                        spec_content, msg_list, result = await task
                        if spec_content:
                            yield f"data: {json.dumps({'type': 'step', 'step': 'generate_spec', 'status': 'success', 'message': 'Spec文档生成成功', 'spec_content': spec_content, 'progress': 85})}\n\n"

                            # 更新 Version 状态为 SPEC_GENERATED
                            await self._update_version_status(
                                version_id=version_id,
                                status=VersionStatus.SPEC_GENERATED,
                                spec_content=spec_content
                            )
                            yield f"data: {json.dumps({'type': 'step', 'step': 'update_version_status', 'status': 'success', 'message': 'Version状态已更新为SPEC_GENERATED', 'progress': 87})}\n\n"
                        else:
                            yield f"data: {json.dumps({'type': 'step', 'step': 'generate_spec', 'status': 'error', 'message': 'Spec文档生成失败', 'spec_content': spec_content, 'progress': 85})}\n\n"
                        # 保存会话

                        try:
                            if msg_list:
                                await message_repo.create_message(
                                    session_id=session_id,
                                    role='assistant',
                                    content="".join(msg_list),
                                    tool_name=None,
                                    tool_input=None,
                                    tool_result=None,
                                )
                            if result:
                                await message_repo.create_message(
                                    session_id=session_id,
                                    role='assistant',
                                    content="".join(result),
                                    tool_name=None,
                                    tool_input=None,
                                    tool_result=None,
                                )
                        except Exception as e:
                            logger.error(f"Failed to save messages: {e}", exc_info=True)

                        yield f"data: {json.dumps({'type': 'step', 'step': 'generate_code', 'status': 'progress', 'message': '正在生成预览供您确认...', 'progress': 90})}\n\n"

                        yield f"data: {json.dumps({'type': 'step', 'step': 'generate_code', 'status': 'success', 'message': f'预览生成成功', 'progress': 100})}\n\n"

                    except Exception as e:
                        logger.error(f"Spec/Code generation failed: {e}", exc_info=True)
                        # await self.module_repo.delete_module(module_id=module_id)
                        yield f"data: {json.dumps({'type': 'error', 'message': f'文档/代码生成失败: {str(e)}'})}\n\n"
                        return

                # 完成
                module = await self.module_repo.get_module_by_id(module_id=module_id)
                response_data = ModuleResponse.model_validate(module)

                yield f"data: {json.dumps({'type': 'complete', 'module_id': module_id, 'session_id': session_id, 'data': response_data.model_dump(mode='json')})}\n\n"

            else:
                # NODE类型：简单创建
                yield f"data: {json.dumps({'type': 'step', 'step': 'create_node', 'status': 'progress', 'message': '创建NODE类型模块...', 'progress': 50})}\n\n"

                menu = {
                    "full_name": data.name,
                    "english_name": data.code,
                    "url_address": '/',
                    "enable_mark": 1,
                    "sort_code": self.sort_code
                }
                insert_id = self.db.insert(table='sys_module', data=menu)
                module_data.update({"url_id": insert_id})


                module = await self.module_repo.create_module(data=module_data, created_by=created_by)

                yield f"data: {json.dumps({'type': 'step', 'step': 'create_node', 'status': 'success', 'message': 'NODE模块创建成功', 'progress': 100})}\n\n"
                yield f"data: {json.dumps({'type': 'complete', 'module_id': module.id, 'data': ModuleResponse.model_validate(module).model_dump(mode='json')})}\n\n"

        except Exception as e:
            logger.error(f"Stream creation failed: {e}", exc_info=True)
            # if module_id:
            #     await self.module_repo.delete_module(module_id=module_id)
            # if insert_id:
            #     self.db.execute_update("DELETE FROM sys_module WHERE id=%s", insert_id)
            yield f"data: {json.dumps({'type': 'error', 'message': f'创建失败: {str(e)}'})}\n\n"

    @log_print
    async def get_module(self, module_id: int):
        """获取模块详情"""
        try:
            module = await self.module_repo.get_module_by_id(module_id=module_id)
            if not module:
                return BaseResponse.not_found(message=f"模块 ID {module_id} 不存在")

            return BaseResponse.success(
                data=ModuleResponse.model_validate(module),
                message="获取模块详情成功"
            )
        except Exception as e:
            logger.error(f"获取模块详情失败: {e}", exc_info=True)
            return BaseResponse.error(message=f"获取模块详情失败: {str(e)}")

    @log_print
    async def get_module_by_session_id(self, session_id: str):
        """根据 session_id 获取模块（仅 POINT 类型）"""
        try:
            module = await self.module_repo.get_module_by_session_id(session_id=session_id)
            if not module:
                return BaseResponse.not_found(message=f"Session ID '{session_id}' 对应的模块不存在")

            return BaseResponse.success(
                data=ModuleResponse.model_validate(module),
                message="获取模块详情成功"
            )
        except Exception as e:
            logger.error(f"获取模块详情失败: {e}", exc_info=True)
            return BaseResponse.error(message=f"获取模块详情失败: {str(e)}")

    @log_print
    async def update_module(
        self,
        module_id: int,
        data: ModuleUpdate,
        updated_by: Optional[str] = None
    ):
        """更新模块信息"""
        try:
            # Check if module exists
            existing = await self.module_repo.get_module_by_id(module_id=module_id)
            if not existing:
                return BaseResponse.not_found(message=f"模块 ID {module_id} 不存在")

            # If updating code, check for duplicates in the same project
            if data.code and data.code != existing.code:
                duplicate = await self.module_repo.get_module_by_code(
                    project_id=existing.project_id,
                    is_active = 1,
                    code=data.code
                )
                if duplicate:
                    return BaseResponse.error(
                        message=f"模块代码 '{data.code}' 在项目中已存在"
                    )

            # Verify new parent if specified
            if data.parent_id is not None:
                if data.parent_id == module_id:
                    return BaseResponse.error(message="模块不能设置自己为父模块")
                parent = await self.module_repo.get_module_by_id(module_id=data.parent_id)
                if not parent:
                    return BaseResponse.not_found(message=f"父模块 ID {data.parent_id} 不存在")
                if parent.project_id != existing.project_id:
                    return BaseResponse.error(message="父模块必须属于同一项目")

            module = await self.module_repo.update_module(
                module_id=module_id,
                data=data,
                updated_by=updated_by
            )
            return BaseResponse.success(
                data=ModuleResponse.model_validate(module),
                message="模块更新成功"
            )
        except Exception as e:
            logger.error(f"更新模块失败: {e}", exc_info=True)
            return BaseResponse.error(message=f"更新模块失败: {str(e)}")

    @log_print
    async def delete_module(self, module_id: int):
        """
        删除模块（递归删除子模块，清理所有相关资源）

        清理内容：
        1. 递归删除所有子模块
        2. 删除 framework 数据库中的 sys_module 记录
        3. 对于 POINT 类型模块：
           - 停止并删除容器
           - 删除工作空间目录
           - 将所有版本记录状态改为 DELETED
        4. 删除模块数据库记录
        """
        try:
            module = await self.module_repo.get_module_by_id(module_id=module_id)
            if not module:
                return BaseResponse.not_found(message=f"模块 ID {module_id} 不存在")

            # 递归删除函数
            async def delete_module_recursive(mod: any) -> None:
                """递归删除模块及其所有子模块"""
                # 1. 先递归删除所有子模块
                children = await self.module_repo.get_children_modules(parent_id=mod.id)
                for child in children:
                    await delete_module_recursive(child)

                logger.info(f"Deleting module: {mod.name} (ID: {mod.id}, Type: {mod.type})")

                # 2. 删除 framework 数据库中的 sys_module 记录
                if mod.url_id:
                    try:
                        self.db.execute_update("DELETE FROM sys_module WHERE id=%s", (mod.url_id,))
                        logger.info(f"Deleted sys_module record: url_id={mod.url_id}")
                    except Exception as e:
                        logger.warning(f"Failed to delete sys_module record: {e}")

                # 3. 处理 POINT 类型模块的特殊资源
                if mod.type == ModuleType.POINT:
                    # 3.1 停止并删除容器
                    if mod.session_id:
                        try:
                            executor = get_sandbox_executor()
                            await executor.stop_container(mod.session_id)
                            logger.info(f"Stopped and deleted container for session: {mod.session_id}")
                        except Exception as e:
                            logger.warning(f"Failed to stop container: {e}")

                    # 3.2 删除工作空间目录
                    if mod.workspace_path:
                        try:
                            import shutil
                            workspace = Path(mod.workspace_path)
                            if workspace.exists():
                                shutil.rmtree(workspace)
                                logger.info(f"Removed workspace: {mod.workspace_path}")
                        except Exception as e:
                            logger.warning(f"Failed to remove workspace: {e}")

                    # 3.3 将所有版本记录状态改为 DELETED
                    try:
                        from app.db.models.version import VersionStatus
                        versions = await self.version_repo.get_versions_by_module(
                            module_id=mod.id,
                            skip=0,
                            limit=1000  # 获取所有版本
                        )

                        for version in versions:
                            if version.status != VersionStatus.DELETED.value:
                                from app.db.schemas import VersionUpdate
                                version_update = VersionUpdate(
                                    status=VersionStatus.DELETED.value
                                )
                                await self.version_repo.update_version(
                                    version_id=version.id,
                                    data=version_update
                                )

                        if versions:
                            logger.info(f"Updated {len(versions)} version records to DELETED status for module {mod.id}")
                    except Exception as e:
                        logger.warning(f"Failed to update version status: {e}")

                # 4. 删除模块数据库记录
                await self.module_repo.delete_module(module_id=mod.id)
                logger.info(f"Deleted module from database: {mod.name} (ID: {mod.id})")

            # 开始递归删除
            await delete_module_recursive(module)

            return BaseResponse.success(message=f"模块 '{module.name}' 及其子模块已成功删除")
        except Exception as e:
            logger.error(f"删除模块失败: {e}", exc_info=True)
            return BaseResponse.error(message=f"删除模块失败: {str(e)}")

    @log_print
    async def pull_module_code(self, module_id: int):
        """拉取 POINT 类型模块的最新代码"""
        try:
            module = await self.module_repo.get_module_by_id(module_id=module_id)
            if not module:
                return BaseResponse.not_found(message=f"模块 ID {module_id} 不存在")

            if module.type != ModuleType.POINT:
                return BaseResponse.error(message="只有 POINT 类型模块支持拉取代码")

            if not module.workspace_path:
                return BaseResponse.error(message="模块没有工作空间路径")

            # Get project info for git config
            project = await self.project_repo.get_project_by_id(project_id=module.project_id)
            if not project:
                return BaseResponse.error(message="关联的项目不存在")

            # Pull latest code
            from git import Repo

            workspace = Path(module.workspace_path)
            if not workspace.exists() or not (workspace / ".git").exists():
                return BaseResponse.error(message="工作空间不存在或不是Git仓库")

            repo = Repo(module.workspace_path)
            origin = repo.remotes.origin
            origin.pull()

            return BaseResponse.success(message="代码拉取成功")
        except Exception as e:
            logger.error(f"拉取代码失败: {e}", exc_info=True)
            return BaseResponse.error(message=f"拉取代码失败: {str(e)}")

    @log_print
    async def restart_module_container(self, module_id: int):
        """重启 POINT 类型模块的容器"""
        try:
            module = await self.module_repo.get_module_by_id(module_id=module_id)
            if not module:
                return BaseResponse.not_found(message=f"模块 ID {module_id} 不存在")

            if module.type != ModuleType.POINT:
                return BaseResponse.error(message="只有 POINT 类型模块有容器")

            if not module.session_id:
                return BaseResponse.error(message="模块没有关联的会话")

            # Restart container
            # executor = get_sandbox_executor()
            # await executor.restart_container(
            #     session_id=module.session_id,
            #     workspace_path=module.workspace_path or ""
            # )

            return BaseResponse.success(message="容器重启成功")
        except Exception as e:
            logger.error(f"重启容器失败: {e}", exc_info=True)
            return BaseResponse.error(message=f"重启容器失败: {str(e)}")

    async def optimize_module_stream(
        self,
        session_id: str,
        content: str,
        updated_by: Optional[str] = None
    ) -> AsyncGenerator[str, None]:
        """
        流式优化已创建的POINT类型模块（SSE）

        根据用户新的需求对现有代码进行优化，实时返回优化进度

        流程：
        1. 查找并验证模块
        2. 加载历史会话上下文
        3. 读取工作空间代码
        4. 结合所有上下文更新spec_content
        5. 生成优化后的代码
        6. 创建commit和version
        7. 重启容器

        生成的事件：
        - connected: 连接建立
        - step: {step, status, message, progress}
        - complete: {module_id, commit_id, version_id}
        - error: {message}
        """
        try:
            # 发送连接确认
            yield f"data: {json.dumps({'type': 'connected', 'message': '开始优化模块'}, ensure_ascii=False)}\n\n"
            await asyncio.sleep(0)

            # 步骤1: 查找模块
            yield f"data: {json.dumps({'type': 'step', 'step': 'find_module', 'status': 'progress', 'message': '查找模块...', 'progress': 5}, ensure_ascii=False)}\n\n"

            module = await self.module_repo.get_module_by_session_id(session_id=session_id)
            if not module:
                yield f"data: {json.dumps({'type': 'error', 'message': f'Session ID {session_id} 对应的模块不存在'}, ensure_ascii=False)}\n\n"
                return

            if module.type != ModuleType.POINT:
                yield f"data: {json.dumps({'type': 'error', 'message': '只能优化POINT类型模块'}, ensure_ascii=False)}\n\n"
                return

            yield f"data: {json.dumps({'type': 'step', 'step': 'find_module', 'status': 'success', 'message': f'找到模块: {module.name}', 'progress': 10}, ensure_ascii=False)}\n\n"

            # 步骤2: 验证工作空间
            yield f"data: {json.dumps({'type': 'step', 'step': 'verify_workspace', 'status': 'progress', 'message': '验证工作空间...', 'progress': 15}, ensure_ascii=False)}\n\n"

            workspace_path = module.workspace_path
            if not workspace_path or not Path(workspace_path).exists():
                yield f"data: {json.dumps({'type': 'error', 'message': '工作空间不存在'}, ensure_ascii=False)}\n\n"
                return

            yield f"data: {json.dumps({'type': 'step', 'step': 'verify_workspace', 'status': 'success', 'message': '工作空间验证成功', 'progress': 20}, ensure_ascii=False)}\n\n"

            await message_repo.create_message(
                session_id=session_id,
                role='user',
                content=content,
                tool_name=None,
                tool_input=None,
                tool_result=None,
            )

            # 步骤5: 生成更新后的spec文档
            yield f"data: {json.dumps({'type': 'step', 'step': 'update_spec', 'status': 'progress', 'message': '根据优化需求更新spec文档...', 'progress': 45}, ensure_ascii=False)}\n\n"


            message_queue = asyncio.Queue()
            try:
                task = asyncio.create_task(generate_code_from_spec(
                    spec_content=module.require_content,
                    workspace_path=workspace_path,
                    session_id=session_id,
                    module_code=module.code,
                    module_name=module.name,
                    module_url=module.url,
                    task_type="",
                    message_queue=message_queue, )
                )
                buffer = ""
                while True:
                    chunk = await message_queue.get()
                    if chunk is None:
                        # 最后 flush 剩余内容（即使没有 \n）
                        if buffer:
                            yield f"data: {json.dumps({'type': 'step', 'step': 'ai_think', 'status': 'success', 'message': 'ai思考...', 'ai_message': buffer, 'progress': 50}, ensure_ascii=False)}\n\n"
                        break
                    # 累加到缓冲区
                    buffer += chunk.content
                    # 持续检查缓冲区中是否有完整的行（以 \n 结尾）
                    while "\n\n" in buffer:
                        # 分割出第一个完整行（包含 \n）
                        line, buffer = buffer.split("\n\n", 1)
                        # 推送这一整行
                        yield f"data: {json.dumps({'type': 'step', 'step': 'ai_think', 'status': 'success', 'message': 'ai思考...', 'ai_message': line, 'progress': 60}, ensure_ascii=False)}\n\n"
                spec_content, msg_list, result = await task
                if spec_content:
                    yield f"data: {json.dumps({'type': 'step', 'step': 'update_spec', 'status': 'success', 'message': 'spec更新成功', 'spec_content': spec_content, 'progress': 80}, ensure_ascii=False)}\n\n"
                else:
                    yield f"data: {json.dumps({'type': 'error', 'message': 'spec更新失败', 'spec_content': spec_content, 'progress': 100}, ensure_ascii=False)}\n\n"
                    return
                # 保存会话
                try:
                    if msg_list:
                        await message_repo.create_message(
                            session_id=session_id,
                            role='assistant',
                            content="".join(msg_list),
                            tool_name=None,
                            tool_input=None,
                            tool_result=None,
                        )
                    if result:
                        await message_repo.create_message(
                            session_id=session_id,
                            role='assistant',
                            content="".join(result),
                            tool_name=None,
                            tool_input=None,
                            tool_result=None,
                        )
                except Exception as e:
                    logger.error(f"Failed to save messages: {e}", exc_info=True)
                module_update = ModuleUpdate(spec_content=spec_content)
                module_repo.update_module(module_id=module.session_id, module_update=module_update)

            except Exception as e:
                logger.error(f"Failed to generate code: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'代码生成失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 完成
            yield f"data: {json.dumps({'type': 'complete', 'module_id': module.id, 'spec_content': spec_content, 'version_id': '', 'message': '优化完成', 'progress': 100}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"Optimization stream failed: {e}", exc_info=True)
            yield f"data: {json.dumps({'type': 'error', 'message': f'优化失败: {str(e)}'}, ensure_ascii=False)}\n\n"

    async def build_module_stream(
            self,
            session_id: str,
            content: str,
            updated_by: Optional[str] = None
    ) -> AsyncGenerator[str, None]:
        try:
            # 发送连接确认
            yield f"data: {json.dumps({'type': 'connected', 'message': '开始生成代码'}, ensure_ascii=False)}\n\n"
            await asyncio.sleep(0)

            # 步骤1: 查找模块
            yield f"data: {json.dumps({'type': 'step', 'step': 'find_module', 'status': 'progress', 'message': '查找模块...', 'progress': 5}, ensure_ascii=False)}\n\n"

            module = await self.module_repo.get_module_by_session_id(session_id=session_id)
            if not module:
                yield f"data: {json.dumps({'type': 'error', 'message': f'Session ID {session_id} 对应的模块不存在'}, ensure_ascii=False)}\n\n"
                return

            if module.type != ModuleType.POINT:
                yield f"data: {json.dumps({'type': 'error', 'message': '只能优化POINT类型模块'}, ensure_ascii=False)}\n\n"
                return

            yield f"data: {json.dumps({'type': 'step', 'step': 'find_module', 'status': 'success', 'message': f'找到模块: {module.name}', 'progress': 10}, ensure_ascii=False)}\n\n"

            # 步骤2: 验证工作空间
            yield f"data: {json.dumps({'type': 'step', 'step': 'verify_workspace', 'status': 'progress', 'message': '验证工作空间...', 'progress': 15}, ensure_ascii=False)}\n\n"

            workspace_path = module.workspace_path
            if not workspace_path or not Path(workspace_path).exists():
                yield f"data: {json.dumps({'type': 'error', 'message': '工作空间不存在'}, ensure_ascii=False)}\n\n"
                return

            yield f"data: {json.dumps({'type': 'step', 'step': 'verify_workspace', 'status': 'success', 'message': '工作空间验证成功', 'progress': 20}, ensure_ascii=False)}\n\n"

            # 步骤2.5: 检查容器限制并创建/更新 Version
            yield f"data: {json.dumps({'type': 'step', 'step': 'check_container_limit', 'status': 'progress', 'message': '检查容器限制...', 'progress': 22}, ensure_ascii=False)}\n\n"

            # 检查容器数量
            can_create, error_msg = await self._check_container_limit()
            if not can_create:
                yield f"data: {json.dumps({'type': 'error', 'message': error_msg}, ensure_ascii=False)}\n\n"
                return

            yield f"data: {json.dumps({'type': 'step', 'step': 'check_container_limit', 'status': 'success', 'message': '容器限制检查通过', 'progress': 24}, ensure_ascii=False)}\n\n"

            # 查找或创建 Version 记录
            yield f"data: {json.dumps({'type': 'step', 'step': 'prepare_version', 'status': 'progress', 'message': '准备版本记录...', 'progress': 26}, ensure_ascii=False)}\n\n"

            # 查找 SPEC_GENERATED 状态的 Version
            version = await self.version_repo.get_version_by_module_and_status(
                module_id=module.id,
                status=VersionStatus.SPEC_GENERATED.value
            )

            if version:
                # 更新现有 Version 状态为 CODE_BUILDING
                await self._update_version_status(
                    version_id=version.id,
                    status=VersionStatus.CODE_BUILDING
                )
                version_id = version.id
                yield f"data: {json.dumps({'type': 'step', 'step': 'prepare_version', 'status': 'success', 'message': f'Version {version_id} 状态更新为CODE_BUILDING', 'progress': 28}, ensure_ascii=False)}\n\n"
            else:
                # 创建新的 Version（状态：CODE_BUILDING）
                version_code = f"v1.0.0-{datetime.now().strftime('%Y%m%d%H%M%S')}"
                version_data = VersionCreate(
                    code=version_code,
                    module_id=module.id,
                    msg="[SpecCoding Auto Commit] - Code building",
                    commit="pending",
                    status=VersionStatus.CODE_BUILDING.value
                )
                version = await self.version_repo.create_version(data=version_data, created_by=updated_by)
                version_id = version.id
                yield f"data: {json.dumps({'type': 'step', 'step': 'prepare_version', 'status': 'success', 'message': f'创建Version {version_id}', 'progress': 28}, ensure_ascii=False)}\n\n"

            yield f"data: {json.dumps({'type': 'step', 'step': 'code_build', 'status': 'success', 'message': '开始生成代码', 'progress': 30}, ensure_ascii=False)}\n\n"
            # 步骤3: 生成代码
            message_queue = asyncio.Queue()
            try:
                task = asyncio.create_task(generate_code_from_spec(
                    spec_content=module.require_content,
                    workspace_path=workspace_path,
                    session_id=session_id,
                    module_code=module.code,
                    module_name=module.name,
                    module_url=module.url,
                    task_type="build",
                    message_queue=message_queue, )
                )
                buffer = ""
                while True:
                    chunk = await message_queue.get()
                    if chunk is None:
                        # 最后 flush 剩余内容（即使没有 \n）
                        if buffer:
                            yield f"data: {json.dumps({'type': 'step', 'step': 'ai_think', 'status': 'success', 'message': 'ai思考...', 'ai_message': buffer, 'progress': 50}, ensure_ascii=False)}\n\n"
                        break
                    # 累加到缓冲区
                    buffer += chunk.content
                    # 持续检查缓冲区中是否有完整的行（以 \n 结尾）
                    while "\n\n" in buffer:
                        # 分割出第一个完整行（包含 \n）
                        line, buffer = buffer.split("\n\n", 1)
                        # 推送这一整行
                        yield f"data: {json.dumps({'type': 'step', 'step': 'ai_think', 'status': 'success', 'message': 'ai思考...', 'ai_message': line, 'progress': 60}, ensure_ascii=False)}\n\n"
                spec_content, msg_list, result = await task
                # 保存会话
                try:
                    if msg_list:
                        await message_repo.create_message(
                            session_id=session_id,
                            role='assistant',
                            content="".join(msg_list),
                            tool_name=None,
                            tool_input=None,
                            tool_result=None,
                        )
                    if result:
                        await message_repo.create_message(
                            session_id=session_id,
                            role='assistant',
                            content="".join(result),
                            tool_name=None,
                            tool_input=None,
                            tool_result=None,
                        )
                except Exception as e:
                    logger.error(f"Failed to save messages: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'step', 'step': 'code_build', 'status': 'success', 'message': '代码生成成功', 'progress': 45}, ensure_ascii=False)}\n\n"

            except Exception as e:
                logger.error(f"Failed to generate code: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'代码生成失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤4: Commit 代码
            yield f"data: {json.dumps({'type': 'step', 'step': 'commit', 'status': 'progress', 'message': '提交代码...', 'progress': 50}, ensure_ascii=False)}\n\n"
            # 使用 GitHubService 进行本地 commit
            commit_id = None
            try:
                # Commit message
                commit_message = f"[SpecCoding Auto Commit] - {module.name} ({module.code}) 功能实现"

                # 执行commit
                from git import Repo
                repo = Repo(workspace_path)

                # Add all changes
                repo.git.add(A=True)
                # Check if there are staged changes to commit
                staged_files = repo.git.diff("--cached", "--name-only")

                # Check if there are changes to commit
                if staged_files.strip():
                    # Commit
                    commit = repo.index.commit(commit_message)
                    commit_id = commit.hexsha[:12]  # 使用前12位
                    repo.git.push()

                    logger.info(f"Code committed successfully: {commit_id}")
                    yield f"data: {json.dumps({'type': 'step', 'step': 'commit', 'status': 'success', 'message': f'commit完成,id: {commit_id}', 'progress': 60}, ensure_ascii=False)}\n\n"
                else:
                    logger.warning("No changes to commit")
                    yield f"data: {json.dumps({'type': 'step', 'step': 'commit', 'status': 'skipped', 'message': '无变更，跳过提交', 'progress': 60}, ensure_ascii=False)}\n\n"
                    return

            except Exception as e:
                logger.error(f"Failed to commit code: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'代码提交失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤5: 更新 Version 状态为 BUILD_COMPLETED
            yield f"data: {json.dumps({'type': 'step', 'step': 'update_version', 'status': 'progress', 'message': '更新版本状态...', 'progress': 70}, ensure_ascii=False)}\n\n"

            try:
                # 更新 Version 的 commit 和状态
                version_update = VersionUpdate(
                    commit=commit_id,
                    status=VersionStatus.BUILD_COMPLETED.value,
                    msg=f"{module.name} 代码构建完成: {content[:100]}",
                    module_id=module.id,
                    spec_content=spec_content,
                )
                await self.version_repo.update_version(
                    version_id=version_id,
                    data=version_update
                )

                yield f"data: {json.dumps({'type': 'step', 'step': 'update_version', 'status': 'success', 'message': 'Version状态已更新为BUILD_COMPLETED', 'progress': 75}, ensure_ascii=False)}\n\n"
            except Exception as e:
                logger.error(f"Failed to update version: {e}")
                yield f"data: {json.dumps({'type': 'step', 'step': 'update_version', 'status': 'warning', 'message': '版本更新失败', 'progress': 75}, ensure_ascii=False)}\n\n"

            # 步骤6: 清理容器
            yield f"data: {json.dumps({'type': 'step', 'step': 'cleanup_container', 'status': 'progress', 'message': '清理容器...', 'progress': 78}, ensure_ascii=False)}\n\n"

            try:
                cleanup_success = await self._cleanup_container(module.id, session_id)
                if cleanup_success:
                    yield f"data: {json.dumps({'type': 'step', 'step': 'cleanup_container', 'status': 'success', 'message': '容器清理成功', 'progress': 80}, ensure_ascii=False)}\n\n"
                else:
                    yield f"data: {json.dumps({'type': 'step', 'step': 'cleanup_container', 'status': 'warning', 'message': '容器清理失败', 'progress': 80}, ensure_ascii=False)}\n\n"
            except Exception as e:
                logger.error(f"Failed to cleanup container: {e}")
                yield f"data: {json.dumps({'type': 'step', 'step': 'cleanup_container', 'status': 'warning', 'message': '容器清理失败', 'progress': 80}, ensure_ascii=False)}\n\n"

            # 步骤5: 更新模块的latest_commit_id
            yield f"data: {json.dumps({'type': 'step', 'step': 'update_module', 'status': 'progress', 'message': '更新模块信息...', 'progress': 90}, ensure_ascii=False)}\n\n"

            try:
                module_update = ModuleUpdate(latest_commit_id=commit_id, spec_content=spec_content)
                await self.module_repo.update_module(
                    module_id=module.id,
                    data=module_update
                )
                yield f"data: {json.dumps({'type': 'step', 'step': 'update_module', 'status': 'success', 'message': '模块信息更新成功', 'progress': 100}, ensure_ascii=False)}\n\n"
            except Exception as e:
                logger.error(f"Failed to update module: {e}")
                yield f"data: {json.dumps({'type': 'step', 'step': 'update_module', 'status': 'warning', 'message': '模块更新失败', 'progress': 100}, ensure_ascii=False)}\n\n"

            # 完成
            yield f"data: {json.dumps({'type': 'complete', 'module_id': module.id, 'spec_content': spec_content, 'version_id': version_id, 'commit_id': commit_id, 'message': '代码构建完成，容器已清理'}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"Optimization stream failed: {e}", exc_info=True)
            yield f"data: {json.dumps({'type': 'error', 'message': f'优化失败: {str(e)}'}, ensure_ascii=False)}\n\n"

    async def _convert_file_to_markdown(self, file: UploadFile, temp_path: Path) -> str:
        """
        将上传的文件转换为Markdown格式

        Args:
            file: 上传的文件对象
            temp_path: 临时文件路径

        Returns:
            转换后的Markdown内容

        Raises:
            ValueError: 不支持的文件格式
        """
        file_ext = temp_path.suffix.lower()

        try:
            if file_ext == '.md':
                # Markdown文件直接读取
                content = temp_path.read_text(encoding='utf-8')
                return content

            elif file_ext == '.txt':
                # 纯文本文件直接读取
                content = temp_path.read_text(encoding='utf-8')
                return content

            elif file_ext == '.docx':
                # Word文档转换为Markdown
                try:
                    from docx import Document
                except ImportError:
                    raise ValueError("需要安装 python-docx 库来处理 .docx 文件。请运行: pip install python-docx")

                doc = Document(temp_path)
                markdown_lines = []

                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        markdown_lines.append("")
                        continue

                    # 根据样式转换为Markdown
                    style_name = para.style.name.lower()

                    if 'heading 1' in style_name:
                        markdown_lines.append(f"# {text}")
                    elif 'heading 2' in style_name:
                        markdown_lines.append(f"## {text}")
                    elif 'heading 3' in style_name:
                        markdown_lines.append(f"### {text}")
                    elif 'heading 4' in style_name:
                        markdown_lines.append(f"#### {text}")
                    elif 'heading 5' in style_name:
                        markdown_lines.append(f"##### {text}")
                    elif 'heading 6' in style_name:
                        markdown_lines.append(f"###### {text}")
                    else:
                        markdown_lines.append(text)

                # 处理表格
                for table in doc.tables:
                    markdown_lines.append("")
                    for i, row in enumerate(table.rows):
                        cells = [cell.text.strip() for cell in row.cells]
                        markdown_lines.append("| " + " | ".join(cells) + " |")
                        if i == 0:
                            # 添加表头分隔符
                            markdown_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
                    markdown_lines.append("")

                return "\n".join(markdown_lines)

            else:
                raise ValueError(f"不支持的文件格式: {file_ext}。支持的格式: .docx, .md, .txt")

        except Exception as e:
            logger.error(f"文件转换失败: {e}", exc_info=True)
            raise

    async def upload_file_and_create_module_stream(
        self,
        file: UploadFile,
        session_id: str
    ) -> AsyncGenerator[str, None]:
        """
        上传文件并流式处理PRD分解任务

        流程：
        1. 接收文件上传
        2. 根据文件类型转换为Markdown
        3. 保存到 workspace/{session_id}/prd.md
        4. 调用 chat_stream 进行 prd-decompose 任务
        5. 读取生成的 FEATURE_TREE.md 和 METADATA.json
        6. 返回给前端

        Args:
            file: 上传的文件
            session_id: 会话ID

        Yields:
            SSE格式的事件消息
        """
        temp_file_path = None

        try:
            # 发送连接确认
            yield f"data: {json.dumps({'type': 'connected'})}\n\n"
            await asyncio.sleep(0)

            # 步骤1: 准备工作空间路径
            yield f"data: {json.dumps({'type': 'step', 'step': 'prepare_workspace', 'status': 'progress', 'message': '准备工作空间...', 'progress': 5}, ensure_ascii=False)}\n\n"

            # 构建workspace路径: settings.workspace_base_path / session_id
            workspace_dir = Path.home() / "workspace" / session_id
            workspace_dir.mkdir(parents=True, exist_ok=True)
            workspace_path = str(workspace_dir.absolute())

            yield f"data: {json.dumps({'type': 'step', 'step': 'prepare_workspace', 'status': 'success', 'message': f'工作空间: {workspace_path}', 'progress': 10}, ensure_ascii=False)}\n\n"

            # 步骤2: 接收并保存文件
            yield f"data: {json.dumps({'type': 'step', 'step': 'upload_file', 'status': 'progress', 'message': '正在接收文件...', 'progress': 15}, ensure_ascii=False)}\n\n"

            # 创建临时文件
            file_ext = Path(file.filename).suffix
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
                temp_file_path = Path(temp_file.name)
                content = await file.read()
                temp_file.write(content)

            yield f"data: {json.dumps({'type': 'step', 'step': 'upload_file', 'status': 'success', 'message': f'文件接收成功: {file.filename}', 'progress': 20}, ensure_ascii=False)}\n\n"

            # 步骤3: 转换文件为Markdown
            yield f"data: {json.dumps({'type': 'step', 'step': 'convert_file', 'status': 'progress', 'message': '正在转换文件格式...', 'progress': 25}, ensure_ascii=False)}\n\n"

            try:
                markdown_content = await self._convert_file_to_markdown(file, temp_file_path)

                # 保存到 workspace/{session_id}/prd.md
                prd_file_path = workspace_dir / "prd.md"

                with open(prd_file_path, "w", encoding="utf-8") as f:
                    f.write(markdown_content)

                yield f"data: {json.dumps({'type': 'step', 'step': 'convert_file', 'status': 'success', 'message': f'文件已保存到: {prd_file_path}', 'progress': 30}, ensure_ascii=False)}\n\n"

            except ValueError as e:
                yield f"data: {json.dumps({'type': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"
                return
            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'message': f'文件转换失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤4: 调用 chat_stream 进行 prd-decompose 任务
            yield f"data: {json.dumps({'type': 'step', 'step': 'prd_decompose', 'status': 'progress', 'message': '开始PRD分解任务...', 'progress': 35}, ensure_ascii=False)}\n\n"

            try:
                # 获取沙箱服务
                # 调用 chat_stream，传入 prd.md 的绝对路径和 task_type
                prompt = str(prd_file_path.absolute())

                logger.info(f"Starting prd-decompose task: session_id={session_id}, prompt={prompt}")

                # 流式处理 prd-decompose 任务
                buffer = ""
                async for chat_msg in self.agent_service.chat_stream(
                    prompt=prompt,
                    session_id=session_id,
                    task_type="prd-decompose",
                ):
                    # 转发 chat 消息作为进度更新
                    msg_dict = chat_msg.to_dict()

                    # 根据消息类型调整进度展示
                    if chat_msg.type in ("text", "text_delta"):
                        # 文本消息，累加到缓冲区，按 \n\n 分隔输出
                        buffer += chat_msg.content
                        while "\n\n" in buffer:
                            line, buffer = buffer.split("\n\n", 1)
                            yield f"data: {json.dumps({'type': 'step', 'step': 'ai_think', 'status': 'success', 'message': 'ai思考...', 'ai_message': line, 'progress': 50}, ensure_ascii=False)}\n\n"
                    elif chat_msg.type == "tool_use":
                        # 工具调用，打印日志
                        tool_name = msg_dict.get('tool_name', 'unknown')
                        logger.info(f"Tool use: {tool_name}")
                    elif chat_msg.type == "error":
                        yield f"data: {json.dumps({'type': 'error', 'message': f'PRD分解失败: {chat_msg.content}'}, ensure_ascii=False)}\n\n"
                        return

                    await asyncio.sleep(0.01)

                # 最后 flush 剩余内容（即使没有 \n\n）
                if buffer:
                    yield f"data: {json.dumps({'type': 'step', 'step': 'ai_think', 'status': 'success', 'message': 'ai思考...', 'ai_message': buffer, 'progress': 50}, ensure_ascii=False)}\n\n"

                yield f"data: {json.dumps({'type': 'step', 'step': 'prd_decompose', 'status': 'success', 'message': 'PRD分解任务完成', 'progress': 80}, ensure_ascii=False)}\n\n"

            except Exception as e:
                logger.error(f"PRD decompose failed: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'PRD分解失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤5: 读取生成的文件
            yield f"data: {json.dumps({'type': 'step', 'step': 'read_results', 'status': 'progress', 'message': '读取生成的文件...', 'progress': 85}, ensure_ascii=False)}\n\n"

            try:
                # 根据文档，生成的文件路径为: {workspace_path}/docs/PRD-GEN/
                # 支持不区分大小写查找 PRD-GEN 目录
                prd_gen_dir = self._find_prd_gen_dir(workspace_dir)
                feature_tree_path = prd_gen_dir / "FEATURE_TREE.md"
                metadata_path = prd_gen_dir / "METADATA.json"

                # 检查文件是否存在
                if not feature_tree_path.exists():
                    yield f"data: {json.dumps({'type': 'error', 'message': f'未找到文件: {feature_tree_path}'}, ensure_ascii=False)}\n\n"
                    return

                if not metadata_path.exists():
                    yield f"data: {json.dumps({'type': 'error', 'message': f'未找到文件: {metadata_path}'}, ensure_ascii=False)}\n\n"
                    return

                # 读取文件内容
                feature_tree_content = feature_tree_path.read_text(encoding='utf-8')
                metadata_content = metadata_path.read_text(encoding='utf-8')

                # 解析 JSON
                try:
                    metadata_json = json.loads(metadata_content)
                except json.JSONDecodeError as e:
                    logger.error(f"Failed to parse METADATA.json: {e}")
                    metadata_json = {"error": "Invalid JSON format"}

                yield f"data: {json.dumps({'type': 'step', 'step': 'read_results', 'status': 'success', 'message': '文件读取成功', 'progress': 95}, ensure_ascii=False)}\n\n"

            except Exception as e:
                logger.error(f"Failed to read generated files: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'读取生成文件失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤6: 创建项目和节点
            data = await self.create_modules_from_metadata(session_id=session_id)
            result_data = {
                'type': 'complete',
                'message': 'PRD处理完成',
                'progress': 100,
                'data': {
                    'feature_tree': feature_tree_content,
                    'metadata': metadata_json,
                    'prd_path': str(prd_file_path.absolute()),
                    'feature_tree_path': str(feature_tree_path.absolute()),
                    'metadata_path': str(metadata_path.absolute()),
                    'project_id': data['project_id']
                }
            }

            yield f"data: {json.dumps(result_data, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"Upload and PRD decompose failed: {e}", exc_info=True)
            yield f"data: {json.dumps({'type': 'error', 'message': f'处理失败: {str(e)}'}, ensure_ascii=False)}\n\n"

        finally:
            # 清理临时文件
            if temp_file_path and temp_file_path.exists():
                try:
                    temp_file_path.unlink()
                    logger.info(f"Temporary file deleted: {temp_file_path}")
                except Exception as e:
                    logger.warning(f"Failed to delete temporary file: {e}")

    async def prd_change_stream(
        self,
        session_id: str,
        selected_content: str,
        msg: str,
    ) -> AsyncGenerator[str, None]:
        """
        PRD 修改任务（流式）

        根据用户反馈修改已有的 PRD 内容

        流程：
        1. 验证 session_id 对应的目录和文件是否存在
        2. 构建 prompt: User Review on "{selected_content}", msg: "{msg}"
        3. 调用 chat_stream 进行 prd-change 任务
        4. 读取更新后的 FEATURE_TREE.md 和 METADATA.json
        5. 返回给前端

        Args:
            session_id: 会话ID（必须与原始PRD的session_id一致）
            selected_content: 选中的内容
            msg: 提出的需求

        Yields:
            SSE格式的事件消息
        """
        try:
            # 发送连接确认
            yield f"data: {json.dumps({'type': 'connected'})}\n\n"
            await asyncio.sleep(0)

            # 步骤1: 验证 session_id 和文件
            yield f"data: {json.dumps({'type': 'step', 'step': 'validate_session', 'status': 'progress', 'message': '验证会话和文件...', 'progress': 5}, ensure_ascii=False)}\n\n"

            # 构建workspace路径
            workspace_dir = Path.home() / "workspace" / session_id

            # 检查目录是否存在
            if not workspace_dir.exists():
                yield f"data: {json.dumps({'type': 'error', 'message': f'会话 {session_id} 对应的工作空间不存在'}, ensure_ascii=False)}\n\n"
                return

            # 检查必要文件是否存在
            prd_file_path = workspace_dir / "prd.md"
            prd_gen_dir = self._find_prd_gen_dir(workspace_dir)
            feature_tree_path = prd_gen_dir / "FEATURE_TREE.md"
            metadata_path = prd_gen_dir / "METADATA.json"

            missing_files = []
            if not prd_file_path.exists():
                missing_files.append("prd.md")
            if not feature_tree_path.exists():
                missing_files.append("FEATURE_TREE.md")
            if not metadata_path.exists():
                missing_files.append("METADATA.json")

            if missing_files:
                message = f"缺少必要文件: {', '.join(missing_files)}"
                yield f"data: {json.dumps({'type': 'error', 'message': message}, ensure_ascii=False)}\n\n"
                return

            workspace_path = str(workspace_dir.absolute())
            yield f"data: {json.dumps({'type': 'step', 'step': 'validate_session', 'status': 'success', 'message': '会话和文件验证成功', 'progress': 10}, ensure_ascii=False)}\n\n"

            # 步骤2: 构建 prompt
            yield f"data: {json.dumps({'type': 'step', 'step': 'build_prompt', 'status': 'progress', 'message': '构建请求...', 'progress': 15}, ensure_ascii=False)}\n\n"

            # 按照规范构建 prompt: User Review on "选中的内容", msg: "提出的需求"
            prompt = f'User Review on "{selected_content}", msg: "{msg}"'

            logger.info(f"PRD change prompt: {prompt}")
            yield f"data: {json.dumps({'type': 'step', 'step': 'build_prompt', 'status': 'success', 'message': 'Prompt 构建完成', 'progress': 20}, ensure_ascii=False)}\n\n"

            # 步骤3: 调用 chat_stream 进行 prd-change 任务
            yield f"data: {json.dumps({'type': 'step', 'step': 'prd_change', 'status': 'progress', 'message': '开始PRD修改任务...', 'progress': 25}, ensure_ascii=False)}\n\n"

            try:

                logger.info(f"Starting prd-change task: session_id={session_id}, prompt={prompt}")

                # 流式处理 prd-change 任务
                async for chat_msg in self.agent_service.chat_stream(
                    prompt=prompt,
                    session_id=session_id,
                    task_type="prd-change",
                ):
                    # 转发 chat 消息作为进度更新
                    msg_dict = chat_msg.to_dict()

                    # 根据消息类型调整进度展示
                    if chat_msg.type in ("text", "text_delta"):
                        # 文本消息，显示为 prd_change 进度（25-75%）
                        yield f"data: {json.dumps({'type': 'step', 'step': 'prd_change', 'status': 'progress', 'message': chat_msg.content[:100], 'progress': 50}, ensure_ascii=False)}\n\n"
                    elif chat_msg.type == "tool_use":
                        # 工具调用
                        tool_name = msg_dict.get('tool_name', 'unknown')
                        yield f"data: {json.dumps({'type': 'step', 'step': 'prd_change', 'status': 'progress', 'message': f'正在执行: {tool_name}', 'progress': 60}, ensure_ascii=False)}\n\n"
                    elif chat_msg.type == "error":
                        yield f"data: {json.dumps({'type': 'error', 'message': f'PRD修改失败: {chat_msg.content}'}, ensure_ascii=False)}\n\n"
                        return

                    await asyncio.sleep(0.01)

                yield f"data: {json.dumps({'type': 'step', 'step': 'prd_change', 'status': 'success', 'message': 'PRD修改任务完成', 'progress': 75}, ensure_ascii=False)}\n\n"

            except Exception as e:
                logger.error(f"PRD change failed: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'PRD修改失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤4: 读取更新后的文件
            yield f"data: {json.dumps({'type': 'step', 'step': 'read_results', 'status': 'progress', 'message': '读取更新后的文件...', 'progress': 85}, ensure_ascii=False)}\n\n"

            try:
                # 重新读取文件内容（已经被更新）
                feature_tree_content = feature_tree_path.read_text(encoding='utf-8')
                metadata_content = metadata_path.read_text(encoding='utf-8')

                # 解析 JSON
                try:
                    metadata_json = json.loads(metadata_content)
                except json.JSONDecodeError as e:
                    logger.error(f"Failed to parse METADATA.json: {e}")
                    metadata_json = {"error": "Invalid JSON format"}

                yield f"data: {json.dumps({'type': 'step', 'step': 'read_results', 'status': 'success', 'message': '文件读取成功', 'progress': 95}, ensure_ascii=False)}\n\n"

            except Exception as e:
                logger.error(f"Failed to read updated files: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'读取更新文件失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤5: 返回结果
            result_data = {
                'type': 'complete',
                'message': 'PRD修改完成',
                'progress': 100,
                'data': {
                    'feature_tree': feature_tree_content,
                    'metadata': metadata_json,
                    'feature_tree_path': str(feature_tree_path.absolute()),
                    'metadata_path': str(metadata_path.absolute()),
                }
            }

            yield f"data: {json.dumps(result_data, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"PRD change stream failed: {e}", exc_info=True)
            yield f"data: {json.dumps({'type': 'error', 'message': f'处理失败: {str(e)}'}, ensure_ascii=False)}\n\n"

    async def confirm_prd_stream(
        self,
        session_id: str,
    ) -> AsyncGenerator[str, None]:
        """
        PRD 审阅确认任务（流式）

        用户已确认PRD修改完成，进行确认

        流程：
        1. 验证 session_id 对应的目录和文件是否存在
        2. 调用 chat_stream 进行 confirm-prd 任务（prompt 为空字符串）
        3. 读取更新后的 FEATURE_TREE.md 和 METADATA.json
        4. 返回给前端

        Args:
            session_id: 会话ID（必须与原始PRD的session_id一致）

        Yields:
            SSE格式的事件消息
        """
        try:
            # 发送连接确认
            yield f"data: {json.dumps({'type': 'connected'})}\n\n"
            await asyncio.sleep(0)

            # 步骤1: 验证 session_id 和文件
            yield f"data: {json.dumps({'type': 'step', 'step': 'validate_session', 'status': 'progress', 'message': '验证会话和文件...', 'progress': 5}, ensure_ascii=False)}\n\n"

            # 构建workspace路径
            workspace_dir = Path.home() / "workspace" / session_id

            # 检查目录是否存在
            if not workspace_dir.exists():
                yield f"data: {json.dumps({'type': 'error', 'message': f'会话 {session_id} 对应的工作空间不存在'}, ensure_ascii=False)}\n\n"
                return

            # 检查必要文件是否存在
            prd_file_path = workspace_dir / "prd.md"
            prd_gen_dir = self._find_prd_gen_dir(workspace_dir)
            feature_tree_path = prd_gen_dir / "FEATURE_TREE.md"
            metadata_path = prd_gen_dir / "METADATA.json"

            missing_files = []
            if not prd_file_path.exists():
                missing_files.append("prd.md")
            if not feature_tree_path.exists():
                missing_files.append("FEATURE_TREE.md")
            if not metadata_path.exists():
                missing_files.append("METADATA.json")

            if missing_files:
                message = f"缺少必要文件: {', '.join(missing_files)}"
                yield f"data: {json.dumps({'type': 'error', 'message': message}, ensure_ascii=False)}\n\n"
                return

            workspace_path = str(workspace_dir.absolute())
            yield f"data: {json.dumps({'type': 'step', 'step': 'validate_session', 'status': 'success', 'message': '会话和文件验证成功', 'progress': 10}, ensure_ascii=False)}\n\n"

            # 步骤2: 调用 chat_stream 进行 confirm-prd 任务
            yield f"data: {json.dumps({'type': 'step', 'step': 'confirm_prd', 'status': 'progress', 'message': '开始PRD确认任务...', 'progress': 15}, ensure_ascii=False)}\n\n"

            try:
                # 获取沙箱服务
                # confirm-prd 任务的 prompt 为空字符串
                prompt = ""

                logger.info(f"Starting confirm-prd task: session_id={session_id}")

                # 流式处理 confirm-prd 任务
                async for chat_msg in self.agent_service.chat_stream(
                    prompt=prompt,
                    session_id=session_id,
                    task_type="confirm-prd",
                ):
                    # 转发 chat 消息作为进度更新
                    msg_dict = chat_msg.to_dict()

                    # 根据消息类型调整进度展示
                    if chat_msg.type in ("text", "text_delta"):
                        # 文本消息，显示为 confirm_prd 进度（15-75%）
                        yield f"data: {json.dumps({'type': 'step', 'step': 'confirm_prd', 'status': 'progress', 'message': chat_msg.content[:100], 'progress': 50}, ensure_ascii=False)}\n\n"
                    elif chat_msg.type == "tool_use":
                        # 工具调用
                        tool_name = msg_dict.get('tool_name', 'unknown')
                        yield f"data: {json.dumps({'type': 'step', 'step': 'confirm_prd', 'status': 'progress', 'message': f'正在执行: {tool_name}', 'progress': 60}, ensure_ascii=False)}\n\n"
                    elif chat_msg.type == "error":
                        yield f"data: {json.dumps({'type': 'error', 'message': f'PRD确认失败: {chat_msg.content}'}, ensure_ascii=False)}\n\n"
                        return

                    await asyncio.sleep(0.01)

                yield f"data: {json.dumps({'type': 'step', 'step': 'confirm_prd', 'status': 'success', 'message': 'PRD确认任务完成', 'progress': 75}, ensure_ascii=False)}\n\n"

            except Exception as e:
                logger.error(f"PRD confirm failed: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'PRD确认失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤3: 读取更新后的文件
            yield f"data: {json.dumps({'type': 'step', 'step': 'read_results', 'status': 'progress', 'message': '读取更新后的文件...', 'progress': 85}, ensure_ascii=False)}\n\n"

            try:
                # 重新读取文件内容（已经被更新）
                feature_tree_content = feature_tree_path.read_text(encoding='utf-8')
                metadata_content = metadata_path.read_text(encoding='utf-8')

                # 解析 JSON
                try:
                    metadata_json = json.loads(metadata_content)
                except json.JSONDecodeError as e:
                    logger.error(f"Failed to parse METADATA.json: {e}")
                    metadata_json = {"error": "Invalid JSON format"}

                yield f"data: {json.dumps({'type': 'step', 'step': 'read_results', 'status': 'success', 'message': '文件读取成功', 'progress': 95}, ensure_ascii=False)}\n\n"

            except Exception as e:
                logger.error(f"Failed to read updated files: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'读取更新文件失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤4: 返回结果
            result_data = {
                'type': 'complete',
                'message': 'PRD确认完成',
                'progress': 100,
                'data': {
                    'feature_tree': feature_tree_content,
                    'metadata': metadata_json,
                    'feature_tree_path': str(feature_tree_path.absolute()),
                    'metadata_path': str(metadata_path.absolute()),
                }
            }

            yield f"data: {json.dumps(result_data, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"PRD confirm stream failed: {e}", exc_info=True)
            yield f"data: {json.dumps({'type': 'error', 'message': f'处理失败: {str(e)}'}, ensure_ascii=False)}\n\n"

    @log_print
    async def create_modules_from_metadata(self, session_id: str):
        """
        根据 METADATA.json 批量创建 project 和 modules

        流程：
        1. 读取 {workspace}/{session_id}/docs/PRD-GEN/METADATA.json
        2. 根据 system_info 创建 project
        3. 根据 features 递归创建 modules
        4. 插入 sys_module 表

        Args:
            session_id: 会话ID

        Returns:
            BaseResponse with created project_id and module count
        """
        try:
            # 步骤1: 读取 METADATA.json
            workspace_dir = Path.home() / "workspace" / session_id
            prd_gen_dir = self._find_prd_gen_dir(workspace_dir)
            metadata_path = prd_gen_dir / "METADATA.json"

            if not metadata_path.exists():
                return BaseResponse.error(message=f"未找到文件: {metadata_path}")

            with open(metadata_path, 'r', encoding='utf-8') as f:
                metadata = json.load(f)

            system_info = metadata.get('system_info', {})
            features = metadata.get('features', [])

            if not system_info:
                return BaseResponse.error(message="METADATA.json 中缺少 system_info")

            # 步骤2: 创建 project
            project_name = system_info.get('name_zh', 'Unnamed Project')
            project_code = system_info.get('name_en', 'Unnamed Project') + '-' + session_id

            # 检查项目是否已存在
            existing_project = await self.project_repo.get_project_by_code(code=project_code)

            if existing_project:
                project_id = existing_project.id
                logger.info(f"Project already exists: {project_name} (id: {project_id})")
            else:
                from app.db.schemas import ProjectCreate
                project_data = ProjectCreate(
                    code=project_code,
                    name=project_name,
                    prd_session_id=session_id,
                )

                project = await self.project_repo.create_project(data=project_data)
                project_id = project.id
                logger.info(f"Created project: {project_name} (id: {project_id})")

            # 步骤3: 递归创建 modules
            created_modules = []
            module_count = 0

            async def create_module_tree(feature_data, parent_id=None, url_parent_id=None, level=0):
                """递归创建模块树"""
                nonlocal module_count

                # 生成模块代码
                module_code = f"{feature_data.get('name_en', f'mod_{uuid.uuid4().hex[:8]}')}"

                # 判断模块类型：is_leaf 为 true 是 POINT，false 是 NODE
                is_leaf = feature_data.get('is_leaf', False)
                module_type = ModuleType.POINT if is_leaf else ModuleType.NODE

                # 创建模块数据
                module_data = ModuleCreate(
                    project_id=project_id,
                    parent_id=parent_id,
                    name=feature_data.get('name_zh', 'Unnamed Module'),
                    code=module_code,
                    type=module_type,
                    url=feature_data.get('url', '/')
                )
                module_data = module_data.model_dump()
                module_data.pop('url_parent_id')

                # 如果是 POINT 类型，生成 session_id 和 workspace_path
                if module_type == ModuleType.POINT:
                    session_id = str(uuid.uuid4())
                    workspace_path = str(settings.workspace_base_path / session_id)

                    # 创建 workspace 目录
                    workspace_dir = Path(workspace_path)
                    workspace_dir.mkdir(parents=True, exist_ok=True)
                    module_data.update({
                        'session_id': session_id,
                        'workspace_path': workspace_path,
                        'branch': 'main',
                        'is_active': 1,
                        'content_status': ContentStatus.PENDING
                    })

                    logger.info(f"POINT module: {module_code}, session_id: {session_id}, workspace: {workspace_path}")

                # 检查模块是否已存在
                existing_module = await self.module_repo.get_module_by_code(
                    project_id=project_id,
                    is_active=1,
                    code=module_code
                )

                if existing_module:
                    logger.info(f"Module already exists: {module_code}")
                    module = existing_module
                    # 获取已存在模块的 url_id
                    current_url_id = module.url_id
                else:
                    # 先插入 sys_module 表，获取 insert_id
                    try:
                        menu = {
                            "full_name": feature_data.get('name_zh', 'Unnamed Module'),
                            "english_name": module_code,
                            "url_address": feature_data.get('url', '/'),
                            "enable_mark": 1,
                            "parent_id": url_parent_id,
                            "sort_code": self.sort_code + level * 100
                        }
                        insert_id = self.db.insert(table='sys_module', data=menu)
                        logger.info(f"Inserted sys_module: {module_code}, url_id: {insert_id}")

                        # 将 url_id 添加到模块数据中
                        module_data.update({'url_id': insert_id})
                        current_url_id = insert_id

                    except Exception as e:
                        logger.error(f"Failed to insert sys_module: {e}")
                        # 如果 sys_module 插入失败，继续创建模块但不设置 url_id
                        current_url_id = None

                    # 创建模块记录
                    module = await self.module_repo.create_module(data=module_data)
                    logger.info(f"Created module: {module.name} (type: {module_type}, code: {module_code}, url_id: {current_url_id})")

                module_info = {
                    'id': module.id,
                    'name': module.name,
                    'code': module.code,
                    'type': module_type.value,
                    'url': module.url,
                    'url_id': current_url_id,
                }

                # 如果是 POINT 类型，添加 session_id 和 workspace_path
                if module_type == ModuleType.POINT:
                    module_info['session_id'] = module.session_id
                    module_info['workspace_path'] = module.workspace_path

                created_modules.append(module_info)
                module_count += 1

                # 递归处理子节点，传递当前模块的 url_id 作为子节点的 parent_id
                children = feature_data.get('children', [])
                for child in children:
                    await create_module_tree(child, parent_id=module.id, url_parent_id=current_url_id, level=level + 1)

            # 创建所有功能模块
            for feature in features:
                await create_module_tree(feature, parent_id=None, url_parent_id=None, level=0)

            return {
                    'project_id': project_id,
                    'project_name': project_name,
                    'module_count': module_count
                }

        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse METADATA.json: {e}", exc_info=True)
            return BaseResponse.error(message=f"JSON 解析失败: {str(e)}")
        except Exception as e:
            logger.error(f"Failed to create modules from metadata: {e}", exc_info=True)
            return BaseResponse.error(message=f"创建失败: {str(e)}")

    async def analyze_prd_module_stream(
        self,
        session_id: str,
        module_name: str,
        prd_session_id: str,
    ) -> AsyncGenerator[str, None]:
        """
        PRD 模块分析任务（流式）

        分析 PRD 中的特定功能模块，生成详细的模块设计文档

        流程：
        1. 验证 session_id 对应的 workspace 和 module
        2. 验证 prd_session_id 对应的 FEATURE_TREE.md 和 prd.md
        3. 构建 prompt: --module "{module_name}" --feature-tree "..." --prd "..."
        4. 调用 chat_stream 进行 analyze-prd 任务
        5. 读取生成的 clarification.md
        6. 保存到 module.require_content
        7. 返回给前端

        Args:
            session_id: 模块的 session_id（每次唯一，避免冲突）
            module_name: 要分析的模块名称
            prd_session_id: PRD 的 session_id（用于定位 FEATURE_TREE.md 和 prd.md）

        Yields:
            SSE格式的事件消息
        """
        try:
            # 发送连接确认
            yield f"data: {json.dumps({'type': 'connected'})}\n\n"
            await asyncio.sleep(0)

            # 步骤1: 验证模块的 session_id 和 workspace
            yield f"data: {json.dumps({'type': 'step', 'step': 'validate_module', 'status': 'progress', 'message': '验证模块信息...', 'progress': 5}, ensure_ascii=False)}\n\n"

            # 查找对应的 module
            module = await self.module_repo.get_module_by_session_id(session_id=session_id)
            if not module:
                yield f"data: {json.dumps({'type': 'error', 'message': f'未找到 session_id 为 {session_id} 的模块'}, ensure_ascii=False)}\n\n"
                return

            yield f"data: {json.dumps({'type': 'step', 'step': 'validate_module', 'status': 'success', 'message': f'模块验证成功: {module.name}', 'progress': 10}, ensure_ascii=False)}\n\n"

            # 步骤2: 验证 PRD 的文件
            yield f"data: {json.dumps({'type': 'step', 'step': 'validate_prd', 'status': 'progress', 'message': '验证PRD文件...', 'progress': 15}, ensure_ascii=False)}\n\n"

            prd_workspace_dir = Path.home() / "workspace" / prd_session_id
            prd_gen_dir = self._find_prd_gen_dir(prd_workspace_dir)
            feature_tree_path = prd_gen_dir / "FEATURE_TREE.md"
            prd_file_path = prd_workspace_dir / "prd.md"

            missing_files = []
            if not feature_tree_path.exists():
                missing_files.append("FEATURE_TREE.md")
            if not prd_file_path.exists():
                missing_files.append("prd.md")

            if missing_files:
                message = f"缺少必要文件: {', '.join(missing_files)}"
                yield f"data: {json.dumps({'type': 'error', 'message': message}, ensure_ascii=False)}\n\n"
                return

            yield f"data: {json.dumps({'type': 'step', 'step': 'validate_prd', 'status': 'success', 'message': 'PRD文件验证成功', 'progress': 20}, ensure_ascii=False)}\n\n"

            # 步骤3: 复制文件到新的 session workspace
            yield f"data: {json.dumps({'type': 'step', 'step': 'copy_files', 'status': 'progress', 'message': '复制PRD文件到新workspace...', 'progress': 22}, ensure_ascii=False)}\n\n"

            import shutil
            import stat
            new_workspace_dir = Path.home() / "workspace" / session_id
            new_workspace_dir.mkdir(parents=True, exist_ok=True)

            try:
                # 复制 FEATURE_TREE.md
                new_feature_tree_path = new_workspace_dir / "FEATURE_TREE.md"
                shutil.copy2(feature_tree_path, new_feature_tree_path)
                # 添加所有用户可读可写权限
                os.chmod(new_feature_tree_path, stat.S_IRUSR | stat.S_IWUSR | stat.S_IRGRP | stat.S_IWGRP | stat.S_IROTH | stat.S_IWOTH)
                logger.info(f"Copied FEATURE_TREE.md to {new_feature_tree_path} with read/write permissions for all users")

                # 复制 prd.md
                new_prd_path = new_workspace_dir / "prd.md"
                shutil.copy2(prd_file_path, new_prd_path)
                # 添加所有用户可读可写权限
                os.chmod(new_prd_path, stat.S_IRUSR | stat.S_IWUSR | stat.S_IRGRP | stat.S_IWGRP | stat.S_IROTH | stat.S_IWOTH)
                logger.info(f"Copied prd.md to {new_prd_path} with read/write permissions for all users")

                # 复制 PRD_OVERVIEW.md（如果存在）
                prd_overview_path = prd_gen_dir / "PRD_OVERVIEW.md"
                if prd_overview_path.exists():
                    new_prd_overview_path = new_workspace_dir / "PRD_OVERVIEW.md"
                    shutil.copy2(prd_overview_path, new_prd_overview_path)
                    # 添加所有用户可读可写权限
                    os.chmod(new_prd_overview_path, stat.S_IRUSR | stat.S_IWUSR | stat.S_IRGRP | stat.S_IWGRP | stat.S_IROTH | stat.S_IWOTH)
                    logger.info(f"Copied PRD_OVERVIEW.md to {new_prd_overview_path} with read/write permissions for all users")
                else:
                    logger.warning(f"PRD_OVERVIEW.md not found in {prd_workspace_dir}, skipping")

                yield f"data: {json.dumps({'type': 'step', 'step': 'copy_files', 'status': 'success', 'message': 'PRD文件复制成功', 'progress': 25}, ensure_ascii=False)}\n\n"

            except Exception as e:
                logger.error(f"Failed to copy files: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'文件复制失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤4: 构建 prompt
            yield f"data: {json.dumps({'type': 'step', 'step': 'build_prompt', 'status': 'progress', 'message': '构建分析请求...', 'progress': 27}, ensure_ascii=False)}\n\n"

            # 按照规范构建 prompt: --module "模块名称" --feature-tree "路径" --prd "路径"
            # 使用新workspace中的文件名（不带路径前缀）
            prompt = f'--module "{module_name}" --feature-tree "FEATURE_TREE.md" --prd "prd.md"'

            logger.info(f"analyze-prd prompt: {prompt}")
            yield f"data: {json.dumps({'type': 'step', 'step': 'build_prompt', 'status': 'success', 'message': 'Prompt 构建完成', 'progress': 30}, ensure_ascii=False)}\n\n"

            # 步骤4: 调用 chat_stream 进行 analyze-prd 任务
            yield f"data: {json.dumps({'type': 'step', 'step': 'analyze_prd', 'status': 'progress', 'message': '开始PRD模块分析...', 'progress': 35}, ensure_ascii=False)}\n\n"
            try:
                module_update = ModuleUpdate(content_status=ContentStatus.IN_PROGRESS)
                await self.module_repo.update_module(
                    module_id=module.id,
                    data=module_update
                )

            except Exception as e:
                logger.error(f"Failed to update module: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'更新module失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            try:
                logger.info(f"Starting analyze-prd task: session_id={session_id}, module={module_name}")

                # 流式处理 analyze-prd 任务
                buffer = ""
                async for chat_msg in self.agent_service.chat_stream(
                    prompt=prompt,
                    session_id=session_id,
                    task_type="analyze-prd",
                ):
                    # 转发 chat 消息作为进度更新
                    msg_dict = chat_msg.to_dict()

                    # 根据消息类型调整进度展示
                    if chat_msg.type in ("text", "text_delta"):
                        # 文本消息，累加到缓冲区，按 \n\n 分隔输出
                        buffer += chat_msg.content
                        while "\n\n" in buffer:
                            line, buffer = buffer.split("\n\n", 1)
                            yield f"data: {json.dumps({'type': 'step', 'step': 'ai_think', 'status': 'success', 'message': 'ai思考...', 'ai_message': line, 'progress': 55}, ensure_ascii=False)}\n\n"
                    elif chat_msg.type == "tool_use":
                        # 工具调用，打印日志
                        tool_name = msg_dict.get('tool_name', 'unknown')
                        logger.info(f"Tool use: {tool_name}")
                    elif chat_msg.type == "error":
                        yield f"data: {json.dumps({'type': 'error', 'message': f'PRD分析失败: {chat_msg.content}'}, ensure_ascii=False)}\n\n"
                        return

                    await asyncio.sleep(0.01)

                # 最后 flush 剩余内容（即使没有 \n\n）
                if buffer:
                    yield f"data: {json.dumps({'type': 'step', 'step': 'ai_think', 'status': 'success', 'message': 'ai思考...', 'ai_message': buffer, 'progress': 55}, ensure_ascii=False)}\n\n"

                yield f"data: {json.dumps({'type': 'step', 'step': 'analyze_prd', 'status': 'success', 'message': 'PRD模块分析完成', 'progress': 75}, ensure_ascii=False)}\n\n"

            except Exception as e:
                logger.error(f"PRD analyze failed: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'PRD分析失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤5: 读取生成的 clarification.md
            yield f"data: {json.dumps({'type': 'step', 'step': 'read_clarification', 'status': 'progress', 'message': '读取生成的文档...', 'progress': 80}, ensure_ascii=False)}\n\n"

            try:
                # 根据文档，生成的文件路径为: {workspace_path}/docs/PRD-GEN/clarification.md
                cla_workspace_dir = Path.home() / "workspace" / session_id
                cla_gen_dir = self._find_prd_gen_dir(cla_workspace_dir)
                clarification_path = cla_gen_dir / "clarification.md"
                logger.info("clarification path: {}".format(clarification_path))

                if not clarification_path.exists():
                    yield f"data: {json.dumps({'type': 'error', 'message': f'未找到文件: {clarification_path}'}, ensure_ascii=False)}\n\n"
                    return

                # 读取文件内容
                clarification_content = clarification_path.read_text(encoding='utf-8')

                yield f"data: {json.dumps({'type': 'step', 'step': 'read_clarification', 'status': 'success', 'message': '文档读取成功', 'progress': 85}, ensure_ascii=False)}\n\n"

            except Exception as e:
                logger.error(f"Failed to read clarification.md: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'读取文档失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤6: 保存到 module.require_content
            yield f"data: {json.dumps({'type': 'step', 'step': 'save_content', 'status': 'progress', 'message': '保存到模块...', 'progress': 90}, ensure_ascii=False)}\n\n"

            try:
                module_update = ModuleUpdate(require_content=clarification_content, content_status=ContentStatus.COMPLETED)
                await self.module_repo.update_module(
                    module_id=module.id,
                    data=module_update
                )

                yield f"data: {json.dumps({'type': 'step', 'step': 'save_content', 'status': 'success', 'message': '内容保存成功', 'progress': 95}, ensure_ascii=False)}\n\n"

            except Exception as e:
                logger.error(f"Failed to update module: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': f'保存失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                return

            # 步骤7: 返回结果
            result_data = {
                'type': 'complete',
                'message': 'PRD模块分析完成',
                'progress': 100,
                'data': {
                    'module_id': module.id,
                    'module_name': module.name,
                    'module_code': module.code,
                    'clarification_content': clarification_content,
                    'clarification_path': str(clarification_path.absolute()),
                }
            }

            yield f"data: {json.dumps(result_data, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"analyze-prd stream failed: {e}", exc_info=True)
            yield f"data: {json.dumps({'type': 'error', 'message': f'处理失败: {str(e)}'}, ensure_ascii=False)}\n\n"

    async def prepare_and_generate_spec_stream(
        self,
        session_id: str,
        content: Optional[str] = None,
        created_by: Optional[str] = None
    ) -> AsyncGenerator[str, None]:
        """
        准备环境并生成 Spec（流式）

        流程：
        1. 根据 session_id 查找模块
        2. 验证 project 是否有 git 地址和 token
        3. 检查工作空间是否有代码
           - 没有：拉取代码 → 检查容器阈值 → 创建容器 → 生成 spec
           - 有：检查容器是否存在
             - 没有：检查容器阈值 → 创建容器 → 生成 spec
             - 有：直接生成 spec

        Args:
            session_id: 模块的 session_id
            content: 模块的 需求内容
            created_by: 创建者

        Yields:
            SSE格式的事件消息
        """
        version_id = None

        try:
            # 发送连接确认
            yield f"data: {json.dumps({'type': 'connected'})}\n\n"
            await asyncio.sleep(0)

            # 步骤1: 查找模块
            yield f"data: {json.dumps({'type': 'step', 'step': 'find_module', 'status': 'progress', 'message': '查找模块...', 'progress': 5}, ensure_ascii=False)}\n\n"

            module = await self.module_repo.get_module_by_session_id(session_id=session_id)
            if not module:
                yield f"data: {json.dumps({'type': 'error', 'message': f'Session ID {session_id} 对应的模块不存在'}, ensure_ascii=False)}\n\n"
                return

            if module.type != ModuleType.POINT:
                yield f"data: {json.dumps({'type': 'error', 'message': '只能为POINT类型模块生成Spec'}, ensure_ascii=False)}\n\n"
                return

            yield f"data: {json.dumps({'type': 'step', 'step': 'find_module', 'status': 'success', 'message': f'找到模块: {module.name}', 'progress': 10}, ensure_ascii=False)}\n\n"

            # 步骤2: 验证 project
            yield f"data: {json.dumps({'type': 'step', 'step': 'validate_project', 'status': 'progress', 'message': '验证项目配置...', 'progress': 15}, ensure_ascii=False)}\n\n"

            project = await self.project_repo.get_project_by_id(project_id=module.project_id)
            if not project:
                yield f"data: {json.dumps({'type': 'error', 'message': '关联的项目不存在'}, ensure_ascii=False)}\n\n"
                return

            if not project.codebase:
                yield f"data: {json.dumps({'type': 'error', 'message': '项目没有配置 Git 地址'}, ensure_ascii=False)}\n\n"
                return

            if not project.token:
                yield f"data: {json.dumps({'type': 'error', 'message': '项目没有配置 Git Token'}, ensure_ascii=False)}\n\n"
                return

            yield f"data: {json.dumps({'type': 'step', 'step': 'validate_project', 'status': 'success', 'message': '项目配置验证成功', 'progress': 20}, ensure_ascii=False)}\n\n"

            # 步骤3: 检查工作空间
            yield f"data: {json.dumps({'type': 'step', 'step': 'check_workspace', 'status': 'progress', 'message': '检查工作空间...', 'progress': 25}, ensure_ascii=False)}\n\n"

            workspace_path = module.workspace_path
            if not workspace_path:
                yield f"data: {json.dumps({'type': 'error', 'message': '模块没有工作空间路径'}, ensure_ascii=False)}\n\n"
                return

            workspace_dir = Path(workspace_path)
            has_code = workspace_dir.exists() and (workspace_dir / ".git").exists()

            if has_code:
                yield f"data: {json.dumps({'type': 'step', 'step': 'check_workspace', 'status': 'success', 'message': '工作空间已有代码', 'progress': 30}, ensure_ascii=False)}\n\n"
            else:
                yield f"data: {json.dumps({'type': 'step', 'step': 'check_workspace', 'status': 'success', 'message': '工作空间无代码，需要拉取', 'progress': 30}, ensure_ascii=False)}\n\n"

                # 步骤4: 拉取代码
                yield f"data: {json.dumps({'type': 'step', 'step': 'clone_repo', 'status': 'progress', 'message': '正在克隆代码仓库...', 'progress': 35}, ensure_ascii=False)}\n\n"

                try:
                    service = GitHubService(token=project.token)
                    await service.clone_repo(
                        repo_url=project.codebase,
                        target_path=workspace_path,
                        branch=module.branch or "main",
                    )

                    yield f"data: {json.dumps({'type': 'step', 'step': 'clone_repo', 'status': 'success', 'message': '代码仓库克隆成功', 'progress': 45}, ensure_ascii=False)}\n\n"

                    # 创建分支
                    yield f"data: {json.dumps({'type': 'step', 'step': 'create_branch', 'status': 'progress', 'message': '创建功能分支...', 'progress': 50}, ensure_ascii=False)}\n\n"

                    branch_name = f"{module.branch or 'main'}-{session_id}"
                    await service.create_branch(repo_path=workspace_path, branch_name=branch_name)
                    await self.session_repo.create_session(
                        session_id=session_id,
                        name=project.code + '-' + module.code,
                        workspace_path=workspace_path,
                        github_repo_url=project.codebase,
                        github_branch=module.branch or "main",
                    )

                    yield f"data: {json.dumps({'type': 'step', 'step': 'create_branch', 'status': 'success', 'message': '功能分支创建成功', 'progress': 55}, ensure_ascii=False)}\n\n"

                except Exception as e:
                    yield f"data: {json.dumps({'type': 'error', 'message': f'代码拉取失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                    return

            # 步骤5: 检查容器
            yield f"data: {json.dumps({'type': 'step', 'step': 'check_container', 'status': 'progress', 'message': '检查容器状态...', 'progress': 60}, ensure_ascii=False)}\n\n"

            executor = get_sandbox_executor()
            container_exists = False
            preview_url = None

            try:
                # 尝试获取容器信息
                container_info = await executor.get_container_status(session_id)
                if container_info and container_info.get('status') == 'running':
                    container_exists = True
                    # 获取 preview_url
                    code_port = container_info.get('code_port')
                    if code_port:
                        preview_url = f"{settings.preview_ip}:{code_port}{module.url}"

                    yield f"data: {json.dumps({'type': 'step', 'step': 'check_container', 'status': 'success', 'message': '容器已存在且运行中', 'preview_url': preview_url, 'progress': 65}, ensure_ascii=False)}\n\n"
            except Exception as e:
                logger.info(f"Container not found or not running: {e}")
                yield f"data: {json.dumps({'type': 'step', 'step': 'check_container', 'status': 'success', 'message': '容器不存在，需要创建', 'progress': 65}, ensure_ascii=False)}\n\n"

            # 步骤6: 如果容器不存在，创建容器
            if not container_exists:
                # 检查容器限制
                yield f"data: {json.dumps({'type': 'step', 'step': 'check_container_limit', 'status': 'progress', 'message': '检查容器限制...', 'progress': 67}, ensure_ascii=False)}\n\n"

                can_create, error_msg = await self._check_container_limit()
                if not can_create:
                    yield f"data: {json.dumps({'type': 'error', 'message': error_msg}, ensure_ascii=False)}\n\n"
                    return

                yield f"data: {json.dumps({'type': 'step', 'step': 'check_container_limit', 'status': 'success', 'message': '容器限制检查通过', 'progress': 69}, ensure_ascii=False)}\n\n"

                # 创建容器
                yield f"data: {json.dumps({'type': 'step', 'step': 'create_container', 'status': 'progress', 'message': '创建沙箱容器...', 'progress': 70}, ensure_ascii=False)}\n\n"

                try:
                    container_info = await executor.create_workspace(
                        session_id=session_id,
                        workspace_path=workspace_path
                    )


                    # 获取 preview_url
                    code_port = container_info.get('code_port')
                    if code_port:
                        preview_url = f"{settings.preview_ip}:{code_port}{module.url}"

                    # 更新模块的 container_id和preview_url
                    module_update = ModuleUpdate(container_id=container_info["id"], preview_url=preview_url)
                    await self.module_repo.update_module(module_id=module.id, data=module_update)

                    container_id_short = container_info["id"][:12]
                    yield f"data: {json.dumps({'type': 'step', 'step': 'create_container', 'status': 'success', 'message': f'容器ID: {container_id_short}', 'preview_url': preview_url, 'progress': 75}, ensure_ascii=False)}\n\n"

                except Exception as e:
                    logger.error(f"Container creation failed: {e}")
                    yield f"data: {json.dumps({'type': 'error', 'message': f'容器创建失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                    return

            # 步骤7: 创建 Version 记录
            if content:
                module.require_content = content
            if module.require_content:
                yield f"data: {json.dumps({'type': 'step', 'step': 'create_version', 'status': 'progress', 'message': '创建版本记录...', 'progress': 77}, ensure_ascii=False)}\n\n"

                # 检查是否已有 SPEC_GENERATING 或 SPEC_GENERATED 状态的 Version
                existing_version = await self.version_repo.get_version_by_module_and_status(
                    module_id=module.id,
                    status=VersionStatus.SPEC_GENERATING.value
                )

                if not existing_version:
                    existing_version = await self.version_repo.get_version_by_module_and_status(
                        module_id=module.id,
                        status=VersionStatus.SPEC_GENERATED.value
                    )

                if existing_version:
                    version_id = existing_version.id
                    # 如果是 SPEC_GENERATED，更新回 SPEC_GENERATING
                    if existing_version.status == VersionStatus.SPEC_GENERATED.value:
                        await self._update_version_status(
                            version_id=version_id,
                            status=VersionStatus.SPEC_GENERATING
                        )
                    yield f"data: {json.dumps({'type': 'step', 'step': 'create_version', 'status': 'success', 'message': f'使用已有Version: {version_id}', 'progress': 79}, ensure_ascii=False)}\n\n"
                else:
                    # 创建新 Version
                    version_code = f"v1.0.0-{datetime.now().strftime('%Y%m%d%H%M%S')}"
                    version_data = VersionCreate(
                        code=version_code,
                        module_id=module.id,
                        msg="[SpecCoding Auto Commit] - Spec generation",
                        status=VersionStatus.SPEC_GENERATING.value
                    )

                    version = await self.version_repo.create_version(data=version_data, created_by=created_by)
                    version_id = version.id

                    yield f"data: {json.dumps({'type': 'step', 'step': 'create_version', 'status': 'success', 'message': f'版本ID: {version_id}', 'progress': 79}, ensure_ascii=False)}\n\n"

            # 步骤8: 生成 Spec
            if module.require_content:
                yield f"data: {json.dumps({'type': 'step', 'step': 'generate_spec', 'status': 'progress', 'message': '正在生成技术规格文档...', 'progress': 80}, ensure_ascii=False)}\n\n"

                message_queue = asyncio.Queue()
                try:
                    task = asyncio.create_task(generate_code_from_spec(
                        spec_content=module.require_content,
                        workspace_path=workspace_path,
                        session_id=session_id,
                        module_code=module.code,
                        module_name=module.name,
                        module_url=module.url,
                        task_type="spec",
                        message_queue=message_queue,
                    ))

                    buffer = ""
                    while True:
                        chunk = await message_queue.get()
                        if chunk is None:
                            if buffer:
                                yield f"data: {json.dumps({'type': 'step', 'step': 'ai_think', 'status': 'success', 'message': 'ai思考...', 'ai_message': buffer, 'progress': 85}, ensure_ascii=False)}\n\n"
                            break

                        buffer += chunk.content
                        while "\n\n" in buffer:
                            line, buffer = buffer.split("\n\n", 1)
                            yield f"data: {json.dumps({'type': 'step', 'step': 'ai_think', 'status': 'success', 'message': 'ai思考...', 'ai_message': line, 'progress': 85}, ensure_ascii=False)}\n\n"

                    spec_content, msg_list, result = await task

                    if spec_content:
                        yield f"data: {json.dumps({'type': 'step', 'step': 'generate_spec', 'status': 'success', 'message': 'Spec文档生成成功', 'spec_content': spec_content, 'progress': 90}, ensure_ascii=False)}\n\n"

                        # 更新模块的 spec_content
                        module_update = ModuleUpdate(spec_content=spec_content)
                        await self.module_repo.update_module(module_id=module.id, data=module_update)

                        # 更新 Version 状态为 SPEC_GENERATED
                        if version_id:
                            await self._update_version_status(
                                version_id=version_id,
                                status=VersionStatus.SPEC_GENERATED,
                                spec_content=spec_content
                            )
                            yield f"data: {json.dumps({'type': 'step', 'step': 'update_version_status', 'status': 'success', 'message': 'Version状态已更新为SPEC_GENERATED', 'progress': 95}, ensure_ascii=False)}\n\n"

                        # 保存消息
                        try:
                            if msg_list:
                                await message_repo.create_message(
                                    session_id=session_id,
                                    role='assistant',
                                    content="".join(msg_list),
                                    tool_name=None,
                                    tool_input=None,
                                    tool_result=None,
                                )
                            if result:
                                await message_repo.create_message(
                                    session_id=session_id,
                                    role='assistant',
                                    content="".join(result),
                                    tool_name=None,
                                    tool_input=None,
                                    tool_result=None,
                                )
                        except Exception as e:
                            logger.error(f"Failed to save messages: {e}", exc_info=True)

                        # 完成
                        yield f"data: {json.dumps({'type': 'complete', 'module_id': module.id, 'session_id': session_id, 'version_id': version_id, 'spec_content': spec_content, 'message': 'Spec生成完成'}, ensure_ascii=False)}\n\n"

                    else:
                        yield f"data: {json.dumps({'type': 'error', 'message': 'Spec文档生成失败'}, ensure_ascii=False)}\n\n"

                except Exception as e:
                    logger.error(f"Spec generation failed: {e}", exc_info=True)
                    yield f"data: {json.dumps({'type': 'error', 'message': f'Spec生成失败: {str(e)}'}, ensure_ascii=False)}\n\n"
            else:
                yield f"data: {json.dumps({'type': 'error', 'message': '模块没有需求内容(require_content)'}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"Prepare and generate spec failed: {e}", exc_info=True)
            yield f"data: {json.dumps({'type': 'error', 'message': f'处理失败: {str(e)}'}, ensure_ascii=False)}\n\n"

